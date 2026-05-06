"""
Patch the FYP document with all corrections and improvements.
Creates a new file: 'fyp document FINAL.docx'
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy
import re

SRC = "fyp document with some format changes and next.docx"
DST = "fyp document FINAL.docx"

# ── Text patches: (old_text, new_text) ──────────────────────────────────────
# Exact substring replacements applied across all paragraphs.
TEXT_PATCHES = [
    # EDIT 1 – Abstract & 1.3.2 & scope: 40 → 39 letters
    ("It can recognize 40 Urdu letters.", "It can recognize 39 Urdu letters."),
    ("This mode can recognize 40 Urdu letters.", "This mode can recognize 39 Urdu letters."),
    ("The first part can recognize 40 Urdu alphabet characters.", "The first part can recognize 39 Urdu alphabet characters."),

    # EDIT 2 – TTS engine (abstract)
    ("Says it out loud using Microsoft Edge Text-to-Speech.",
     "Says it out loud using platform-optimized text-to-speech synthesis (Google Android TTS on mobile devices)."),

    # EDIT 2b – TTS engine (section 1.3.3)
    ("Microsoft Edge Text-, to-Speech. This means the system can use a voice that sounds like a person speaking Urdu.",
     "A platform-native text-to-speech engine (Google Android TTS). This means the system can use a voice that sounds natural when speaking Urdu."),

    # EDIT 2c – TTS engine (section 3.3)
    ("Microsoft Edge Text-to-Speech to read out the translated Urdu sentences.",
     "Platform-native text-to-speech (Google Android TTS on mobile, with fallback to system TTS on other platforms) to read out the translated Urdu sentences."),

    # EDIT 2d – TTS engine (section 3.2 state machine)
    ("called Microsoft Edge Text-, to-Speech. This means the system can use a voice that sounds like a person speaking Urdu.",
     "a platform-native text-to-speech engine (Google Android TTS). This means the system can use a voice that sounds natural when speaking Urdu."),

    # EDIT 3 – section 1.3.1 (background interference)
    ("The Listen system ignores things that can interfere with the picture like the background or lighting. This helps the Listen system track the hands accurately.",
     "The Listen system is designed to be robust to various environmental factors. While extreme lighting conditions or severe hand occlusion can temporarily affect performance, the system reliably tracks hands across a wide range of real-world backgrounds and lighting conditions."),

    # EDIT 4 – state machine four → five states (section 1.3.3)
    ("This state machine has four parts: it can be idle, buffering, predicting or committed.",
     "This state machine has five states: idle (waiting for hands), signing (capturing and buffering hand motion), predicting (running inference every 3 frames via EMA-smoothed confidence), committed (locking and outputting the recognized word), and cooldown (preventing duplicate recognition for 0.8 seconds)."),

    # EDIT 5 – stride frames: 5 → 3
    ("triggered only once every 5 frames,", "triggered only once every 3 frames,"),

    # EDIT 6 – early stopping epoch 26 → 41
    ("when the validation loss reached an inflection point, which occurred after the 26th epoch, the training process was stopped.",
     "the training process ran for 41 epochs total, with early stopping configured with a patience of 15 epochs monitoring validation accuracy. Model weights from the epoch with the highest validation accuracy were restored automatically."),

    # EDIT 7 – scope does not include mobile app (remove incorrect claim)
    ("There are some things that the project will not include such as recognizing signs, from people at the same time or automatically correcting grammar mistakes or making a mobile app. These things will be worked on in the future.",
     "There are some things that the project will not include in its current version such as simultaneous multi-signer recognition or automatic grammar correction. A Flutter-based cross-platform mobile application has been implemented as part of the project and connects to the backend server via WebSocket for real-time inference. Further expansion of vocabulary and grammar post-processing are planned for future work."),

    # EDIT 8 – abstract: handsre → hands are (typo)
    ("how the handsre moving", "how the hands are moving"),

    # EDIT 9 – abstract: of how far → regardless of how far
    ("designed to be used by anyone of how far away they are from the camera",
     "designed to be used by anyone regardless of how far they are from the camera"),
]

# ── Abbreviations: full replacement block ────────────────────────────────────
OLD_ABBREV_MARKERS = [
    "cp\tspecific heat capacity",
    "LIST OF SYMBOLS / ABBREVIATIONS",
]

NEW_ABBREVIATIONS = """LIST OF SYMBOLS / ABBREVIATIONS

PSL\t\t\tPakistan Sign Language
ASL\t\t\tAmerican Sign Language
TFLite\t\t\tTensorFlow Lite
CNN\t\t\tConvolutional Neural Network
Conv1D\t\t\tOne-Dimensional Convolutional Layer
LSTM\t\t\tLong Short-Term Memory
BiLSTM\t\t\tBidirectional Long Short-Term Memory
MLP\t\t\tMultilayer Perceptron
FSM\t\t\tFinite State Machine
EMA\t\t\tExponential Moving Average
ReLU\t\t\tRectified Linear Unit (activation function)
TTS\t\t\tText-to-Speech
FPS\t\t\tFrames Per Second
RTL\t\t\tRight-to-Left (Urdu text rendering direction)
API\t\t\tApplication Programming Interface
CPU\t\t\tCentral Processing Unit
GPU\t\t\tGraphics Processing Unit
F\t\t\tFeature dimension (126-D for word model; 42-D for alphabet model)
T\t\t\tTemporal window size (60 frames = 3 seconds at 20 FPS)
IoU\t\t\tIntersection over Union
Adam\t\t\tAdaptive Moment Estimation (optimizer)
L2\t\t\tL2 regularization (weight decay)
UAlpha40\t\tUrdu Alphabet dataset with 40-class sign images (39 used)
MediaPipe\t\tGoogle's cross-platform ML framework for body/hand tracking
WebSocket\t\tFull-duplex real-time communication protocol
FastAPI\t\t\tModern asynchronous Python web framework
Flutter\t\t\tGoogle's cross-platform mobile/desktop UI framework
BSCS\t\t\tBachelor of Science in Computer Science"""


# ── New sections to insert ───────────────────────────────────────────────────
TECH_STACK_TABLE = """
Technology Stack and Dependencies

The Listen system is built on the following core technologies:

Component\t\tTechnology
Hand Tracking\t\tMediaPipe Holistic (CPU-optimised)
Sequence Modelling\tTensorFlow Lite (.tflite)
Backend Server\t\tFastAPI + WebSocket (Python)
Mobile Frontend\t\tFlutter (Dart) — iOS & Android
Text-to-Speech\t\tPlatform-native (Google Android TTS / iOS AVSpeechSynthesizer)
ML Framework\t\tTensorFlow 2.x + Keras
Training Tools\t\tscikit-learn, joblib, NumPy
"""

PERFORMANCE_TABLE = """
Real-Time Inference Performance

Metric\t\t\t\tValue
Word Recognition Accuracy\t98.16%
Top-5 Accuracy\t\t\t99.39%
Supported Word Classes\t\t64
Supported Alphabet Classes\t39
Temporal Window\t\t\t60 frames (3 seconds @ 20 FPS)
Inference Latency (CPU)\t\t< 50 ms end-to-end
Frame Processing Rate\t\t20 FPS
Confidence Threshold\t\t0.70 (70%)
EMA Smoothing Factor\t\t0.60
Inference Stride\t\t\tEvery 3rd frame
"""

IMPL_ARCHITECTURE = """
Implementation Architecture

The system is deployed across three distinct layers:

Backend Server: A FastAPI WebSocket server handles real-time frame-by-frame communication. Two specialised session managers — SignSession (word-level) and AlphabetSession (character-level) — each maintain their own finite state machine and rolling frame buffer per connected client. TFLite interpreters are loaded once at startup and shared across sessions using a threading lock, minimising memory overhead.

Mobile Frontend: A Flutter application provides cross-platform support for iOS and Android. The app streams compressed JPEG frames over WebSocket to the server, renders Urdu text right-to-left, displays real-time confidence scores, and synthesises speech using the platform's native TTS engine.

Model Format: Both classifiers are exported as TensorFlow Lite (.tflite) models. This eliminates mandatory internet connectivity during inference, ensures user privacy, and reduces model size from ~6 MB (Keras .h5) to ~2 MB (TFLite).
"""

UX_FLOW = """
User Interaction Flow (Inference State Machine)

The five-state finite state machine governs the full interaction lifecycle:

1. IDLE — System monitors for hand presence. No processing overhead; buffer is clear.
2. SIGNING — Motion detected; system buffers normalised 126-D landmark vectors. TFLite inference runs every 3rd frame with EMA smoothing applied to class probabilities.
3. PREDICTING — Confidence tracked continuously. Display shows current top prediction and confidence percentage.
4. COMMITTED — When smoothed confidence exceeds 0.70 for a non-idle class, the word is locked. Urdu text appended to sentence; TTS synthesis triggered asynchronously.
5. COOLDOWN — 0.8-second window prevents duplicate recognition. Buffer cleared; system returns to IDLE automatically.

This design requires no button presses and mirrors the natural pause between signed words.
"""

FUTURE_WORK_EXTRA = """
Additional directions for future work include:

- Expanding vocabulary from 64 to 200+ words using active learning and community-sourced data collection
- Implementing grammar correction and semantic post-processing to produce more natural Urdu sentence output
- Adding support for simultaneous multi-signer recognition to facilitate group conversations
- Training a sentence-level continuous sign language recognition model that does not require segmentation between words
- Developing a fully offline mobile app mode that runs the TFLite model directly on-device without a backend server
- Creating a browser-based interface using WebAssembly for access without app installation
- Publishing the Dynamic Pakistan Sign Language Dataset as an open benchmark for the research community
"""


def apply_text_patches(doc):
    for para in doc.paragraphs:
        for (old, new) in TEXT_PATCHES:
            if old in para.text:
                # We do a simple full-run replacement on the paragraph's XML text.
                for run in para.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, new)
                        break
                else:
                    # old spans multiple runs — rebuild the paragraph text
                    full = para.text
                    if old in full:
                        new_full = full.replace(old, new)
                        # Clear all runs and put it in the first run
                        for i, run in enumerate(para.runs):
                            run.text = new_full if i == 0 else ""


def replace_abbreviations_section(doc):
    """Find the LIST OF SYMBOLS paragraph and replace the block up to the next chapter."""
    start_idx = None
    end_idx = None
    for i, para in enumerate(doc.paragraphs):
        if "LIST OF SYMBOLS" in para.text and start_idx is None:
            start_idx = i
        if start_idx is not None and i > start_idx:
            if para.text.strip() in ("LIST OF APPENDICES", "INTRODUCTION", "CHAPTERS"):
                end_idx = i
                break

    if start_idx is None:
        print("  [WARN] LIST OF SYMBOLS section not found — skipping")
        return

    # Replace all paragraphs in that range with our new content
    # We'll clear them and set the first to our new text
    for i in range(start_idx, end_idx if end_idx else start_idx + 30):
        if i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            for run in para.runs:
                run.text = ""
            if i == start_idx:
                if para.runs:
                    para.runs[0].text = NEW_ABBREVIATIONS

    print(f"  Replaced abbreviations section (paragraphs {start_idx}–{end_idx})")


def append_section(doc, heading_text, body_text, heading_level=2):
    """Append a new headed section at the end of the document body."""
    h = doc.add_heading(heading_text, level=heading_level)
    for line in body_text.strip().split("\n"):
        doc.add_paragraph(line)


def main():
    print(f"Opening: {SRC}")
    doc = Document(SRC)

    print("Applying text patches...")
    apply_text_patches(doc)

    print("Replacing abbreviations section...")
    replace_abbreviations_section(doc)

    print("Appending supplementary sections...")
    doc.add_page_break()
    append_section(doc,
                   "SUPPLEMENTARY MATERIAL: SYSTEM DETAILS",
                   TECH_STACK_TABLE + "\n" + PERFORMANCE_TABLE,
                   heading_level=1)
    append_section(doc,
                   "Implementation Architecture",
                   IMPL_ARCHITECTURE,
                   heading_level=2)
    append_section(doc,
                   "User Interaction Flow (Inference State Machine)",
                   UX_FLOW,
                   heading_level=2)
    append_section(doc,
                   "Extended Future Work",
                   FUTURE_WORK_EXTRA,
                   heading_level=2)

    print(f"Saving: {DST}")
    doc.save(DST)
    print("Done.")


if __name__ == "__main__":
    main()
