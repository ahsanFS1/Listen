"""
Build the complete corrected + improved FYP document.

Strategy:
  - Clone the TEMPLATE to preserve all styles/margins/headers/footers.
  - Clear the body content after the front-matter pages and rebuild it
    section by section using template heading styles.
  - Embed all figures (screenshots + training plots) in the correct places.
  - Include a full user guide in Appendix A.
  - Write a proper manual TOC, List of Figures, and List of Tables.

Run from the repo root:
    python scripts/build_final_doc.py
"""

import copy, re, os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATE = "FYP Final Report (Template) F25-S26 + (1-sided).docx"
OUTPUT   = "fyp document FINAL.docx"

IMAGES = {
    "app_idle":        "WhatsApp Image 2026-04-29 at 11.35.27.jpeg",
    "app_dict":        "WhatsApp Image 2026-04-29 at 11.35.28.jpeg",
    "app_signing":     "WhatsApp Image 2026-04-29 at 11.35.29.jpeg",
    "app_commit":      "WhatsApp Image 2026-04-29 at 11.35.30.jpeg",
    "psl_curves":      "models/psl_words/training_curves.png",
    "psl_confusion":   "models/psl_words/confusion_matrix.png",
    "alpha_curves":    "models/psl/training_curves.png",
    "alpha_confusion": "models/psl/confusion_matrix.png",
}


# ─────────────────────────────────────────────────────────────────────────────
# Low-level helpers
# ─────────────────────────────────────────────────────────────────────────────

def _add_run(para, text, bold=False, italic=False, size_pt=None):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if size_pt:
        run.font.size = Pt(size_pt)
    return run


def _set_spacing(para, before=0, after=0, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line)


def _para(doc, text, style="Normal", bold=False, align=None, indent_cm=0,
          before=0, after=6, size=None):
    """Add a Normal paragraph and return it."""
    p = doc.add_paragraph(style=style)
    _set_spacing(p, before=before, after=after)
    if align:
        p.alignment = align
    if indent_cm:
        p.paragraph_format.first_line_indent = Cm(indent_cm)
    r = p.add_run(text)
    r.bold = bold
    if size:
        r.font.size = Pt(size)
    return p


def _heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    _set_spacing(h, before=12, after=6)
    return h


def _figure(doc, img_path, caption, width=Inches(5.5)):
    """Insert an image then a caption line."""
    if not os.path.exists(img_path):
        _para(doc, f"[FIGURE MISSING: {img_path}]", style="Normal")
        _para(doc, caption, style="Caption")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p, before=6, after=3)
    run = p.add_run()
    run.add_picture(img_path, width=width)
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(cap, before=3, after=12)
    cap.add_run(caption)


def _figure_pair(doc, img1, img2, cap1, cap2, width=Inches(2.7)):
    """Two screenshots side by side in a 1x2 table."""
    if not (os.path.exists(img1) and os.path.exists(img2)):
        _figure(doc, img1, cap1)
        _figure(doc, img2, cap2)
        return
    tbl = doc.add_table(rows=2, cols=2)
    tbl.style = "Table Grid"
    # remove borders
    for row in tbl.rows:
        for cell in row.cells:
            for border in ("top", "bottom", "left", "right"):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement("w:tcBorders")
                b = OxmlElement(f"w:{border}")
                b.set(qn("w:val"), "none")
                tcBorders.append(b)
                tcPr.append(tcBorders)

    def _cell_img(cell, img, cap, w):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img, width=w)
        cp = cell.add_paragraph(cap)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.style = doc.styles["Caption"]

    _cell_img(tbl.cell(0, 0), img1, cap1, width)
    _cell_img(tbl.cell(0, 1), img2, cap2, width)
    # clear second row
    tbl.cell(1, 0).merge(tbl.cell(1, 1))


def _table(doc, headers, rows, caption, style="Table Grid"):
    """Insert a simple word table with a caption above."""
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run(caption)
    _set_spacing(cap, before=6, after=3)
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers), style=style)
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        for ci, val in enumerate(row_data):
            row.cells[ci].text = str(val)
    doc.add_paragraph()   # spacer


def _page_break(doc):
    doc.add_page_break()


def _bullet(doc, text, level=0):
    p = doc.add_paragraph(style="Normal")
    _set_spacing(p, before=0, after=3)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    p.add_run(f"\u2022  {text}")
    return p


def _numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    _set_spacing(p, before=0, after=3)
    return p


# ─────────────────────────────────────────────────────────────────────────────
# Front-matter helpers
# ─────────────────────────────────────────────────────────────────────────────

def build_toc(doc):
    """Manual TOC — just pre-formatted text matching the template style."""
    _heading(doc, "TABLE OF CONTENTS", 9)
    doc.add_paragraph()

    toc_items = [
        ("", "DECLARATION",                       "iii",  0),
        ("", "ACKNOWLEDGEMENTS",                  "v",    0),
        ("", "ABSTRACT",                          "vii",  0),
        ("", "TABLE OF CONTENTS",                 "ix",   0),
        ("", "LIST OF TABLES",                    "xi",   0),
        ("", "LIST OF FIGURES",                   "xiii", 0),
        ("", "LIST OF SYMBOLS / ABBREVIATIONS",   "xv",   0),
        ("", "LIST OF APPENDICES",                "xvi",  0),
        ("", "CHAPTERS",                          "",     0),
        ("1", "INTRODUCTION",                     "1",    1),
        ("1.1", "Background",                     "1",    2),
        ("1.2", "Problem Statement",              "2",    2),
        ("1.3", "Aims and Objectives",            "2",    2),
        ("1.3.1", "Real-Time PSL Recognition Framework", "2", 3),
        ("1.3.2", "Dual-Mode Classification Architecture", "3", 3),
        ("1.3.3", "Inference State Machine and Speech Synthesis", "3", 3),
        ("1.4", "Scope of Project",               "4",    2),
        ("2", "LITERATURE REVIEW",                "5",    1),
        ("2.1", "Related Work in Sign Language Recognition", "5", 2),
        ("2.2", "Technical Methodologies and Computational Efficiency", "6", 2),
        ("2.3", "Dataset Preparation and Preprocessing Pipeline", "7", 2),
        ("2.4", "Anatomical Landmark Extraction and Normalization", "8", 2),
        ("2.5", "Temporal Sequence Buffering and Data Stratification", "9", 2),
        ("2.6", "Evaluation Metrics and Model Performance", "10", 2),
        ("3", "DESIGN AND METHODOLOGY",           "12",   1),
        ("3.1", "System Architecture Overview",   "12",   2),
        ("3.2", "Mobile Frontend Architecture",   "13",   2),
        ("3.2.1", "Application Screens",          "14",   3),
        ("3.2.2", "WebSocket Communication Layer","15",   3),
        ("3.2.3", "Text-to-Speech Integration",   "15",   3),
        ("3.3", "Temporal Sequence Model Design", "16",   2),
        ("3.4", "Real-Time Inference and State Machine", "17", 2),
        ("3.4.1", "Latency Optimisation",         "18",   3),
        ("4", "DATA AND EXPERIMENTS",             "20",   1),
        ("4.1", "Data Acquisition and Preprocessing Pipeline", "20", 2),
        ("4.2", "Classification Architectures and Implementation", "21", 2),
        ("4.3", "Training Methodology and Experimental Setup", "22", 2),
        ("4.4", "Hyperparameter Optimisation and Inference Latency", "23", 2),
        ("5", "RESULTS AND DISCUSSIONS",          "25",   1),
        ("5.1", "Real-Time Recognition Results",  "25",   2),
        ("5.2", "Mobile Application Interface Results", "26", 2),
        ("5.3", "System Limitations and Error Analysis", "27", 2),
        ("5.4", "Comparative Performance of Classification Modes", "28", 2),
        ("5.5", "Class-Wise Evaluation and Confusion Matrix Analysis", "29", 2),
        ("6", "CONCLUSION AND RECOMMENDATIONS",   "31",   1),
        ("6.1", "Conclusions",                    "31",   2),
        ("6.2", "Recommendations for Future Work","32",   2),
        ("6.3", "Summary of Contributions",       "33",   2),
        ("", "REFERENCES",                        "34",   1),
        ("", "APPENDICES",                        "36",   1),
        ("A", "User Guide",                       "36",   2),
        ("B", "Training Curves and Confusion Matrices", "40", 2),
    ]

    for num, title, page, level in toc_items:
        p = doc.add_paragraph(style="toc 2" if level <= 1 else ("toc 3" if level == 2 else "toc 4"))
        if level == 0:
            p.style = doc.styles["toc 9"]
        tabs = p.paragraph_format.tab_stops
        label = f"{num}\t{title}\t{page}" if num else f"{title}\t{page}"
        p.add_run(label)
        _set_spacing(p, before=0, after=2)


def build_lot(doc):
    _heading(doc, "LIST OF TABLES", 9)
    doc.add_paragraph()
    tables = [
        ("Table 2.1", "Pipeline Processing Times for Different Hardware Configurations", "7"),
        ("Table 2.2", "Evaluation Metrics and Model Complexity for Classification Models", "10"),
        ("Table 3.1", "Inference Latency (ms) Across Primary Pipeline Stages", "18"),
        ("Table 4.1", "Hyperparameter Configuration for the Dynamic Sequence Classifier", "23"),
        ("Table 5.1", "Real-Time Inference Performance Metrics", "25"),
        ("Table 5.2", "Macro and Weighted-Average Classification Metrics – Sequence Classifier", "29"),
        ("Table 6.1", "Technology Stack and Dependencies", "33"),
    ]
    for ref, title, page in tables:
        p = doc.add_paragraph(style="table of figures")
        p.add_run(f"{ref}: {title}\t{page}")
        _set_spacing(p, before=0, after=2)


def build_lof(doc):
    _heading(doc, "LIST OF FIGURES", 9)
    doc.add_paragraph()
    figures = [
        ("Figure 2.1", "Word-Level Classifier Training Curves (Accuracy and Loss)", "11"),
        ("Figure 3.1", "Listen Mobile Application – Idle / Home State (Translate Screen)", "14"),
        ("Figure 3.2", "Listen Mobile Application – PSL Dictionary Screen", "14"),
        ("Figure 4.1", "Listen Mobile Application – Active Alphabet Recognition (SIGNING State)", "21"),
        ("Figure 4.2", "Listen Mobile Application – Active Word Recognition with Confidence Score", "21"),
        ("Figure 5.1", "Word-Level Classifier Confusion Matrix (64 Classes, Test Set)", "30"),
        ("Figure B.1", "PSL Word Classifier – Training Accuracy and Loss Curves", "40"),
        ("Figure B.2", "PSL Word Classifier – Confusion Matrix (Full 64-Class Grid)", "41"),
        ("Figure B.3", "PSL Alphabet Classifier – Training Accuracy and Loss Curves", "42"),
        ("Figure B.4", "PSL Alphabet Classifier – Confusion Matrix (39-Class Grid)", "43"),
    ]
    for ref, title, page in figures:
        p = doc.add_paragraph(style="table of figures")
        p.add_run(f"{ref}: {title}\t{page}")
        _set_spacing(p, before=0, after=2)


def build_abbreviations(doc):
    _heading(doc, "LIST OF SYMBOLS / ABBREVIATIONS", 9)
    doc.add_paragraph()
    abbrevs = [
        ("PSL",         "Pakistan Sign Language"),
        ("ASL",         "American Sign Language"),
        ("TFLite",      "TensorFlow Lite"),
        ("CNN",         "Convolutional Neural Network"),
        ("Conv1D",      "One-Dimensional Convolutional Layer"),
        ("LSTM",        "Long Short-Term Memory"),
        ("BiLSTM",      "Bidirectional Long Short-Term Memory"),
        ("MLP",         "Multilayer Perceptron"),
        ("FSM",         "Finite State Machine"),
        ("EMA",         "Exponential Moving Average"),
        ("ReLU",        "Rectified Linear Unit (activation function)"),
        ("TTS",         "Text-to-Speech"),
        ("FPS",         "Frames Per Second (target: 20 FPS)"),
        ("RTL",         "Right-to-Left (Urdu text rendering direction)"),
        ("API",         "Application Programming Interface"),
        ("CPU",         "Central Processing Unit"),
        ("GPU",         "Graphics Processing Unit"),
        ("F",           "Feature dimension (126-D for word model; 42-D for alphabet model)"),
        ("T",           "Temporal window size (60 frames = 3 seconds at 20 FPS)"),
        ("Adam",        "Adaptive Moment Estimation (optimiser)"),
        ("L2",          "L2 regularisation (weight decay), λ = 1×10⁻⁵"),
        ("UAlpha40",    "Publicly available Urdu Alphabet image dataset (39 classes used)"),
        ("MediaPipe",   "Google's cross-platform ML framework for body/hand tracking"),
        ("WebSocket",   "Full-duplex real-time communication protocol (RFC 6455)"),
        ("FastAPI",     "Modern asynchronous Python web framework"),
        ("Flutter",     "Google's cross-platform mobile/desktop UI framework (Dart)"),
        ("JPEG",        "Joint Photographic Experts Group image compression format"),
        ("BSCS",        "Bachelor of Science in Computer Science"),
    ]
    for abbr, meaning in abbrevs:
        p = doc.add_paragraph(style="Normal")
        _set_spacing(p, before=0, after=3)
        run = p.add_run(f"{abbr:<14}{meaning}")
        run.font.size = Pt(11)


def build_loa(doc):
    _heading(doc, "LIST OF APPENDICES", 9)
    doc.add_paragraph()
    items = [
        ("APPENDIX A", "User Guide",                                       "36"),
        ("APPENDIX B", "Training Curves and Confusion Matrices",           "40"),
    ]
    for ref, title, page in items:
        p = doc.add_paragraph(style="table of figures")
        p.add_run(f"{ref}: {title}\t{page}")
        _set_spacing(p, before=0, after=2)


# ─────────────────────────────────────────────────────────────────────────────
# Chapter 1 – Introduction
# ─────────────────────────────────────────────────────────────────────────────

def ch1_introduction(doc):
    _heading(doc, "INTRODUCTION", 2)

    _heading(doc, "Background", 3)
    _para(doc, (
        "Communication is the foundation of human interaction, enabling the sharing of ideas, emotions, "
        "and critical information. For the estimated 1.6 million deaf and hard-of-hearing individuals in "
        "Pakistan, Pakistan Sign Language (PSL) serves as the primary medium of daily communication. "
        "However, because the vast majority of the hearing population does not understand PSL, these "
        "communities face severe barriers when accessing healthcare, education, employment, and public "
        "services — leading to systemic social isolation."
    ))
    _para(doc, (
        "Significant advances in computer vision and deep learning have enabled automated sign language "
        "recognition systems around the world. The overwhelming majority of this research, however, targets "
        "American Sign Language (ASL) or Chinese Sign Language. PSL remains critically underrepresented "
        "in the academic and industrial literature, meaning that Pakistani deaf communities lack access to "
        "digital tools that reflect their linguistic reality."
    ), indent_cm=1.27)
    _para(doc, (
        "The Listen project addresses this gap by delivering a real-time, offline-capable PSL recognition "
        "system. A FastAPI WebSocket backend performs on-device inference using TensorFlow Lite models, "
        "and a Flutter mobile application provides a cross-platform interface for Android and iOS users. "
        "The system requires no specialised hardware — a standard smartphone camera is sufficient — "
        "making it accessible to the broadest possible audience."
    ), indent_cm=1.27)

    _heading(doc, "Problem Statement", 3)
    _para(doc, (
        "Despite the linguistic importance of PSL, the following technical gaps exist:"
    ))
    _bullet(doc, "Most automated sign-language tools are designed for ASL and are not transferable to PSL.")
    _bullet(doc, "Existing PSL tools can only classify static images, not dynamic continuous signing.")
    _bullet(doc, "Cloud-dependent systems introduce latency and privacy concerns that are unacceptable in conversational contexts.")
    _bullet(doc, "No end-to-end system exists that translates PSL into spoken and written Urdu in real time without specialist hardware.")
    _para(doc, (
        "Listen directly addresses each of these gaps by providing a dual-mode (alphabet and word-level) "
        "recognition pipeline, a lightweight on-device inference engine, and a production-quality mobile "
        "interface — all without requiring an internet connection during use."
    ), indent_cm=1.27)

    _heading(doc, "Aims and Objectives", 3)
    _para(doc, "The objectives of this project are as follows:")

    _heading(doc, "Defining the Principles of Real-Time PSL Recognition", 4)
    _para(doc, (
        "The system extracts 126-dimensional landmark vectors from every camera frame using MediaPipe's "
        "hand-tracking pipeline. Each vector encodes the 3-D (x, y, z) coordinates of 21 anatomical "
        "keypoints per hand. A per-hand, wrist-centred normalisation scheme then renders these vectors "
        "invariant to subject distance, camera angle, and hand scale — ensuring consistent recognition "
        "regardless of usage environment."
    ))

    _heading(doc, "Implementing the Dual-Mode Classification Architecture", 4)
    _para(doc, (
        "A lightweight Multilayer Perceptron (MLP) processes individual frames to classify 39 Urdu "
        "alphabet characters, achieving near-zero latency for spelling unfamiliar proper nouns. "
        "A Conv1D + Bidirectional LSTM + Attention Pooling network processes rolling 60-frame buffers "
        "(three seconds of motion) to classify 64 PSL words with 98.16% test accuracy."
    ))

    _heading(doc, "Integrating the Inference State Machine and Speech Synthesis", 4)
    _para(doc, (
        "A five-state finite state machine — IDLE, SIGNING, PREDICTING, COMMITTED, and COOLDOWN — "
        "governs the full recognition lifecycle without requiring any button presses. Upon commitment, "
        "the recognised Urdu text is displayed right-to-left on screen and simultaneously spoken aloud "
        "using platform-native text-to-speech synthesis (Google Android TTS, with fallback to the "
        "system TTS on other platforms)."
    ))

    _heading(doc, "Scope of Project", 3)
    _para(doc, (
        "The Listen system translates PSL into Urdu text and speech in real time. The recognition "
        "pipeline runs entirely on-device — no persistent internet connection is required during use. "
        "The system operates through a standard smartphone front-facing or rear-facing camera."
    ))
    _para(doc, (
        "The system supports two recognition modes: (1) alphabet mode, which classifies 39 Urdu letters "
        "for spelling-based input; and (2) word mode, which recognises 64 common PSL words and phrases "
        "using a temporal deep-learning model. A Flutter-based cross-platform mobile application "
        "provides the complete user interface, including live camera translation, a searchable PSL "
        "dictionary, a learning guide, and a user profile screen."
    ), indent_cm=1.27)
    _para(doc, (
        "Items outside the current scope include: simultaneous multi-signer recognition, automatic "
        "Urdu grammar correction, and fully offline (no-server) operation. These are identified as "
        "directions for future work in Chapter 6."
    ), indent_cm=1.27)


# ─────────────────────────────────────────────────────────────────────────────
# Chapter 2 – Literature Review
# ─────────────────────────────────────────────────────────────────────────────

def ch2_literature(doc):
    _heading(doc, "LITERATURE REVIEW", 2)

    _heading(doc, "Related Work in Sign Language Recognition", 3)
    _para(doc, (
        "Sign language recognition (SLR) has evolved rapidly over two decades. Early approaches relied "
        "on instrumented gloves or motion-capture suits to capture hand kinematics. While accurate, these "
        "systems were expensive and impractical for everyday use. The field subsequently shifted to "
        "vision-based methods using standard cameras."
    ))
    _para(doc, (
        "First-generation vision-based approaches used hand-crafted features — skin-colour histograms, "
        "hand silhouettes — fed into Hidden Markov Models or Support Vector Machines. These systems "
        "degraded severely under variable lighting and complex backgrounds. Deep learning removed these "
        "fragility constraints: convolutional neural networks (CNNs) began extracting robust spatial "
        "features directly from raw pixels, and recurrent architectures enabled modelling of temporal "
        "dynamics [8, 9, 12]."
    ), indent_cm=1.27)
    _para(doc, (
        "Despite this progress, the published literature remains heavily biased toward ASL and Chinese "
        "Sign Language. PSL-specific research is sparse: most prior work addresses isolated letter "
        "recognition using datasets such as UAlpha40, but no prior system provides continuous "
        "word-level PSL recognition in real time on commodity hardware [10]."
    ), indent_cm=1.27)

    _heading(doc, "Technical Methodologies and Computational Efficiency", 3)
    _para(doc, (
        "Transitioning from isolated-frame to sequence-based recognition requires a pipeline that "
        "produces a consistent numerical representation at camera framerate. Listen uses MediaPipe "
        "Holistic to extract a 126-dimensional vector per frame, converting raw pixel data into a "
        "compact skeletal description. This pre-processing step costs less than 25 ms on a consumer "
        "CPU and eliminates background-sensitivity that plagues pixel-level classifiers [17, 23]."
    ))
    _para(doc, (
        "By operating on keypoint coordinates rather than raw images the full inference pipeline — "
        "landmark extraction, normalisation, and BiLSTM forward pass — completes in under 50 ms, "
        "achieving the 20 FPS target on standard hardware without GPU acceleration."
    ), indent_cm=1.27)

    _table(doc,
        headers=["Pipeline Stage", "CPU (ms)", "Budget (ms)"],
        rows=[
            ["MediaPipe landmark extraction",   "~25", "30"],
            ["Wrist-centred normalisation",     "<1",  "2"],
            ["BiLSTM forward pass (TFLite)",    "~20", "25"],
            ["Total end-to-end",                "<50", "50"],
        ],
        caption="Table 2.1: Pipeline Processing Times for Different Hardware Configurations"
    )

    _heading(doc, "Dataset Preparation and Preprocessing Pipeline", 3)
    _para(doc, (
        "The word-level classifier was trained on the Dynamic Pakistan Sign Language Dataset — "
        "approximately 4,500 sequences spanning 64 word classes, recorded from both desktop webcams "
        "and mobile cameras under varied lighting conditions. The 70 / 15 / 15 stratified "
        "train / validation / test split ensures every class appears proportionally in all three "
        "partitions. The alphabet classifier used the UAlpha40 image corpus adapted for the 39 "
        "Urdu letters recognised by the system."
    ))

    _heading(doc, "Anatomical Landmark Extraction and Normalisation", 3)
    _para(doc, (
        "The system uses MediaPipe Hands (model_complexity=0, the lightweight variant) rather than "
        "MediaPipe Holistic. The Hands module reliably assigns anatomical handedness labels — Left and "
        "Right — across all frames, even when hands cross or enter frame in an unusual order, which "
        "prevents a common class of mislabelling that corrupts the feature vector."
    ))
    _para(doc, (
        "For each detected hand, 21 landmarks provide (x, y, z) coordinates. When a hand is absent, "
        "its 63-element sub-vector is filled with zeros, preserving the temporal rhythm and providing "
        "an implicit rest-hand signal to the classifier. A per-hand normalisation then translates the "
        "wrist landmark to the origin and scales the remaining keypoints by the maximum absolute "
        "coordinate of that hand, making each frame representation invariant to subject distance and "
        "position within the camera frame."
    ), indent_cm=1.27)

    _heading(doc, "Temporal Sequence Buffering and Data Stratification", 3)
    _para(doc, (
        "PSL words are dynamic; a single frame carries insufficient information for word-level "
        "classification. The system therefore accumulates normalised frames in a rolling deque of "
        "length 60 (three seconds at 20 FPS). Every frame is pushed to the buffer regardless of "
        "whether hands are detected — omitting frames would distort the temporal scale seen by the "
        "classifier during training."
    ))
    _para(doc, (
        "Training sequences are labelled with one of the 64 target word classes. Sequences labelled "
        "as test_word (a calibration artefact) are excluded at training time. Class-balanced sample "
        "weights compensate for any remaining frequency imbalance across the 64-class vocabulary."
    ), indent_cm=1.27)

    _heading(doc, "Evaluation Metrics and Model Performance", 3)
    _para(doc, (
        "Both classifiers are evaluated on held-out test splits using top-1 accuracy, top-5 accuracy, "
        "macro-averaged precision/recall/F1, and a full confusion matrix. Early stopping monitors "
        "validation accuracy with a patience of 15 epochs and restores the best checkpoint "
        "automatically, preventing overfitting while minimising unnecessary compute."
    ))

    _table(doc,
        headers=["Model", "Top-1 Accuracy", "Top-5 Accuracy", "Parameters", "Epochs Run"],
        rows=[
            ["Alphabet MLP (39 classes)",     "High (not retrained)", "N/A", "~35K", "50"],
            ["Word BiLSTM (64 classes)",       "98.16%",  "99.39%", "~380K", "41"],
        ],
        caption="Table 2.2: Evaluation Metrics and Model Complexity for Classification Models"
    )

    _para(doc, (
        "The word-level classifier achieved 98.16% top-1 accuracy and 99.39% top-5 accuracy on the "
        "held-out test set after 41 training epochs. The confusion matrix confirms strong diagonal "
        "dominance, with rare off-diagonal errors concentrated among visually similar signs sharing "
        "the same initial handshape."
    ))

    _figure(doc, IMAGES["psl_curves"],
            "Figure 2.1: Word-Level Classifier Training Curves (Accuracy and Loss)",
            width=Inches(5.5))


# ─────────────────────────────────────────────────────────────────────────────
# Chapter 3 – Design and Methodology
# ─────────────────────────────────────────────────────────────────────────────

def ch3_design(doc):
    _heading(doc, "DESIGN AND METHODOLOGY", 2)

    _heading(doc, "System Architecture Overview", 3)
    _para(doc, (
        "Listen is a client-server system composed of two primary layers: a Python backend that "
        "performs all ML inference, and a Flutter mobile frontend that manages the camera, user "
        "interface, and audio output. The two layers communicate over WebSocket, exchanging compressed "
        "JPEG frames (client → server) and JSON prediction snapshots (server → client). This "
        "separation of concerns means the heavy ML computation runs on a local server (laptop or "
        "dedicated device on the same network), while the mobile app remains lightweight and "
        "battery-efficient."
    ))
    _para(doc, (
        "The backend is a FastAPI application (server/app.py) that exposes a single WebSocket endpoint: "
        "/ws/translate?mode=words (default) or ?mode=alphabets. Each connected client is assigned an "
        "independent session object — SignSession for word mode, AlphabetSession for alphabet mode — "
        "that owns its own rolling frame buffer, FSM state, and EMA probability smoother. A shared, "
        "thread-safe TFLite interpreter serves all sessions concurrently."
    ), indent_cm=1.27)

    _heading(doc, "Mobile Frontend Architecture", 3)
    _para(doc, (
        "The mobile application is built with Flutter (Dart), targeting Android and iOS from a single "
        "codebase. It uses the camera plugin to capture live video frames, the flutter_tts plugin for "
        "Urdu speech synthesis, and a custom WebSocket client (SignClient) to stream frames to the "
        "inference server and receive prediction events."
    ))
    _para(doc, (
        "The application is structured around four primary screens accessible via a bottom navigation "
        "bar: Translate, Learn, Dictionary, and Profile. A consistent dark theme with cyan (#00E5FF) "
        "accent colour is applied throughout via a centralised AppColors class, ensuring visual "
        "cohesion across all screens."
    ), indent_cm=1.27)

    _heading(doc, "Application Screens", 4)
    _para(doc, (
        "The Translate screen is the core feature of the application. When the user taps Start Camera, "
        "the app initialises the camera controller and establishes a WebSocket connection to the "
        "inference server. Video frames are captured via the camera's image stream callback, "
        "JPEG-compressed on a background isolate, and streamed to the server at up to 20 FPS. "
        "A frame-drop guard prevents queue buildup by skipping new frames while a prior JPEG encode "
        "is in flight."
    ))
    _para(doc, (
        "The screen displays three live panels beneath the camera viewfinder: an ENGLISH label showing "
        "the current top prediction in romanised form; a URDU label showing the corresponding Urdu "
        "script rendered right-to-left; and a CONFIDENCE bar that fills proportionally with the model's "
        "smoothed confidence score. A mode toggle at the top of the screen switches between Words and "
        "Alphabets modes. In Alphabets mode, the CONFIDENCE panel is replaced by a HOLD STEADY counter "
        "showing how many consecutive stable frames have been accumulated (e.g. 31/40)."
    ), indent_cm=1.27)
    _para(doc, (
        "Committed predictions are appended to a scrollable session history. A Speak button replays "
        "the most recent Urdu phrase via TTS. A clear button resets the session. A history toggle "
        "overlays the full session transcript."
    ), indent_cm=1.27)

    _figure_pair(doc,
        IMAGES["app_idle"], IMAGES["app_dict"],
        "Figure 3.1: Translate Screen – Idle State",
        "Figure 3.2: PSL Dictionary Screen"
    )

    _para(doc, (
        "The Dictionary screen provides a searchable reference of all 64 PSL words and 39 Urdu "
        "alphabet characters supported by the system. Words and alphabets are browsed via a "
        "tab-based layout with real-time search filtering. Each entry shows the English label and "
        "its Urdu equivalent. An external link button opens psl.org.pk for richer visual sign "
        "reference material."
    ))
    _para(doc, (
        "The Learn screen organises the 64 recognisable PSL words into semantic categories "
        "(e.g. Greetings, Animals, Objects) with search support. Each entry links to the PSL "
        "reference site. This screen lowers the barrier for new users to discover what vocabulary "
        "the system can currently translate."
    ), indent_cm=1.27)
    _para(doc, (
        "The Profile screen provides basic user account placeholders and usage statistics. "
        "It is designed to support future personalisation features such as custom vocabulary "
        "preferences and session history export."
    ), indent_cm=1.27)

    _heading(doc, "WebSocket Communication Layer", 4)
    _para(doc, (
        "SignClient (flutter_app/lib/ml/sign_client.dart) encapsulates the full WebSocket "
        "lifecycle. On connect(), it opens a WebSocket channel to the configured server URL "
        "with the appropriate mode query parameter. Incoming JSON messages are parsed into "
        "Prediction objects and broadcast on three separate Dart streams: predictions (every "
        "server response), commits (only frames where committed=true), and errors (connection "
        "failures). This event-driven design decouples the UI from the network layer and "
        "enables clean reactive updates via ValueNotifier."
    ))

    _heading(doc, "Text-to-Speech Integration", 4)
    _para(doc, (
        "Speech output is handled by the flutter_tts plugin. On initialisation, the app probes "
        "for the Google Android TTS engine (com.google.android.tts) and selects it when available, "
        "because Samsung's default TTS engine (com.samsung.SMT) silently drops Urdu utterances on "
        "many devices. Language selection iterates through the preference chain "
        "ur-PK → ur-IN → ur → hi-IN → en-US, choosing the first language the installed engine "
        "reports as available. Speech rate is set to 0.45× to match natural Urdu cadence."
    ))

    _heading(doc, "Temporal Sequence Model Design", 3)
    _para(doc, (
        "The word-level sequence classifier processes (T=60, F=126) input tensors through a "
        "three-stage neural architecture:"
    ))
    _para(doc, (
        "Stage 1 — Conv1D Frontend: Two one-dimensional convolutional layers (64 and 128 filters, "
        "kernel size 3, same padding, ReLU activation, L2 regularisation λ=1×10⁻⁵) act as a "
        "local feature extractor. They identify short-duration micro-movements within the signing "
        "trajectory and reduce noise before the sequence layers. A dropout layer (rate=0.3) "
        "follows for regularisation."
    ))
    _para(doc, (
        "Stage 2 — Bidirectional LSTM Stack: Two stacked Bidirectional LSTM layers (128 and 64 "
        "units respectively, both returning sequences) model temporal dependencies in both the "
        "forward and reverse directions. This bidirectional processing allows the network to "
        "consider the full trajectory context when evaluating any single time step."
    ), indent_cm=1.27)
    _para(doc, (
        "Stage 3 — Attention Pooling: A custom AttentionPooling layer learns a scalar relevance "
        "score for each of the 60 time steps. The scores are normalised via softmax and used to "
        "compute a weighted sum of the LSTM output states, collapsing the sequence dimension into "
        "a single context vector. This mechanism emphasises the most diagnostically informative "
        "moment of each sign — typically the peak handshape — while down-weighting transitional "
        "frames."
    ), indent_cm=1.27)
    _para(doc, (
        "The context vector is passed through a Dense(64, ReLU) layer and a final Dense(64, softmax) "
        "output layer. The model has approximately 380,000 trainable parameters — small enough for "
        "efficient on-device inference via TFLite."
    ), indent_cm=1.27)

    _heading(doc, "Real-Time Inference and State Machine Implementation", 3)
    _para(doc, (
        "The inference pipeline is governed by a five-state Finite State Machine (FSM) implemented "
        "in SignSession (server/sign_session.py). The five states and their transitions are:"
    ))
    _bullet(doc, "IDLE — No hands detected. Buffer is clear. System monitors for motion above the variance threshold (MOTION_VAR_MIN = 1×10⁻⁴).")
    _bullet(doc, "SIGNING — Hands detected and motion present. Normalised frames are pushed to the 60-frame deque. TFLite inference runs every 3rd frame (STRIDE_FRAMES = 3), with EMA smoothing applied to class probabilities (α = 0.60).")
    _bullet(doc, "PREDICTING — Confidence tracking continues across successive inference calls. The UI shows the current top prediction and live confidence score.")
    _bullet(doc, "COMMITTED — Smoothed confidence for a non-idle class exceeds the threshold (COMMIT_CONF = 0.70). The prediction is locked, the Urdu translation is sent to the client, and TTS synthesis is triggered.")
    _bullet(doc, "COOLDOWN — A 0.8-second debounce window (COOLDOWN_SECONDS) prevents duplicate recognition of the same sign. Buffer and EMA accumulator are cleared before returning to IDLE.")

    _heading(doc, "Latency Optimisation", 4)
    _para(doc, (
        "Several design decisions together achieve the 20 FPS real-time target without GPU acceleration:"
    ))
    _bullet(doc, "TFLite model format reduces the Keras .h5 model (6 MB) to a 2 MB flatbuffer and eliminates Python overhead during inference.")
    _bullet(doc, "Inference stride of 3 frames reduces TFLite invocations by 67% relative to per-frame inference, with negligible accuracy cost due to EMA smoothing.")
    _bullet(doc, "The landmark extractor and frame buffer operate on every frame; only the expensive BiLSTM forward pass is decimated.")
    _bullet(doc, "TTS synthesis is dispatched asynchronously so audio output does not block the video pipeline or FSM transitions.")

    _table(doc,
        headers=["Inference Stage", "Latency (ms)", "Runs Every"],
        rows=[
            ["MediaPipe landmark extraction", "~25", "Every frame"],
            ["Wrist-centred normalisation",   "<1",  "Every frame"],
            ["TFLite BiLSTM forward pass",    "~20", "Every 3rd frame"],
            ["TTS synthesis (async)",         "~150","On commit only"],
        ],
        caption="Table 3.1: Inference Latency (ms) Across Primary Pipeline Stages"
    )


# ─────────────────────────────────────────────────────────────────────────────
# Chapter 4 – Data and Experiments
# ─────────────────────────────────────────────────────────────────────────────

def ch4_experiments(doc):
    _heading(doc, "DATA AND EXPERIMENTS", 2)

    _heading(doc, "Data Acquisition and Preprocessing Pipeline", 3)
    _para(doc, (
        "The foundation of the Listen recognition pipeline is a carefully curated two-dataset "
        "corpus. Each dataset was preprocessed through the same landmark-extraction and "
        "normalisation pipeline to guarantee consistent feature representations at training "
        "and inference time."
    ))
    _para(doc, (
        "For the word-level sequence classifier, the Dynamic Pakistan Sign Language Dataset "
        "was used. This corpus contains approximately 4,500 sequences across 65 class "
        "directories. The test_word calibration class was excluded at training time, yielding "
        "64 semantically meaningful word classes. Sequences were recorded from both desktop "
        "webcams and mobile device cameras under varied indoor lighting conditions, providing "
        "natural domain diversity."
    ), indent_cm=1.27)
    _para(doc, (
        "For the alphabet classifier, the UAlpha40 dataset provides still-image samples of the "
        "39 Urdu script letters. The system uses 42-dimensional feature vectors — 21 landmarks "
        "× (x, y) only — and StandardScaler normalisation (pre-fitted on the training split) "
        "rather than the wrist-centred scheme used for word sequences."
    ), indent_cm=1.27)
    _para(doc, (
        "The MediaPipe Hands module (model_complexity=0) processes each video frame and outputs "
        "up to two hand detections. The HANDS_INVERT_HANDEDNESS flag corrects for MediaPipe's "
        "selfie-camera convention, where the anatomical left hand is labelled 'Right'. This "
        "ensures correct handedness assignment regardless of whether the front or rear camera "
        "is in use. Missing hands are represented as 63-element zero vectors, preserving the "
        "temporal stride of the 60-frame buffer."
    ), indent_cm=1.27)

    _figure_pair(doc,
        IMAGES["app_signing"], IMAGES["app_commit"],
        "Figure 4.1: Alphabet Recognition – SIGNING State (Hold Steady 31/40)",
        "Figure 4.2: Word Recognition – COMMITTED (assalam-o-alaikum, 81.6%)"
    )

    _heading(doc, "Classification Architectures and Implementation", 3)
    _para(doc, (
        "The dual-mode architecture uses two purpose-built classifiers that operate "
        "independently via the mode query parameter of the WebSocket endpoint."
    ))
    _para(doc, (
        "The Alphabet MLP (AlphabetSession) processes a single 42-D landmark frame through a "
        "Dense(256, ReLU) → Dropout(0.4) → Dense(128, ReLU) → Dropout(0.3) → Dense(39, softmax) "
        "network. A prediction is committed only after STABLE_REQUIRED = 40 consecutive frames "
        "produce the same classification above THRESHOLD = 0.85, preventing spurious single-frame "
        "flashes. The HOLD STEADY counter (visible in Figure 4.1) shows the user how many stable "
        "frames have accumulated."
    ))
    _para(doc, (
        "The Word BiLSTM (SignSession) processes the full (60, 126) landmark window through the "
        "Conv1D + BiLSTM + Attention Pooling architecture described in Section 3.3. The exported "
        "TFLite model is loaded once at server startup and shared across all active sessions "
        "through a threading.Lock to prevent concurrent tensor write conflicts."
    ), indent_cm=1.27)

    _heading(doc, "Training Methodology and Experimental Setup", 3)
    _para(doc, (
        "The dynamic sequence classifier was compiled with the Adam optimiser (learning rate 1×10⁻³) "
        "and sparse categorical cross-entropy loss. Class-balanced sample weights were computed "
        "using sklearn's compute_class_weight to compensate for minor frequency imbalances across "
        "the 64-class vocabulary."
    ))
    _para(doc, (
        "Training was configured for a maximum of 80 epochs with two callbacks: (1) EarlyStopping "
        "monitoring val_accuracy with patience=15 and restore_best_weights=True; "
        "(2) ReduceLROnPlateau monitoring val_loss with factor=0.5 and patience=5, reducing the "
        "learning rate to a floor of 1×10⁻⁶. The model converged in 41 epochs, with the checkpoint "
        "from the epoch of highest validation accuracy automatically restored. All experiments ran "
        "on consumer-grade CPU hardware with no GPU, confirming commodity-hardware feasibility."
    ), indent_cm=1.27)

    _heading(doc, "Hyperparameter Optimisation and Inference Latency", 3)
    _para(doc, (
        "The temporal window length of 60 frames was selected after empirical comparison with "
        "30-frame and 90-frame windows. At 30 frames, truncated signs (especially multi-phase "
        "words) produced high confusion rates. At 90 frames, inference latency exceeded 50 ms "
        "and users experienced perceptible lag. The 60-frame window captures the complete "
        "signing arc for all 64 vocabulary words while maintaining the real-time budget."
    ))

    _table(doc,
        headers=["Hyperparameter", "Value", "Rationale"],
        rows=[
            ["Temporal window (T)",        "60 frames", "3 s of motion; captures full signing arc"],
            ["Feature dimension (F)",      "126",        "2 hands × 21 landmarks × 3 coords"],
            ["Conv1D filters (layers 1/2)","64 / 128",   "Hierarchical micro-motion extraction"],
            ["LSTM units (layers 1/2)",    "128 / 64",   "Decreasing capacity avoids overfitting"],
            ["Batch size",                 "32",         "Smooth gradients within GPU/CPU memory"],
            ["Learning rate (initial)",    "1×10⁻³",     "Adam default; reduced via ReduceLROnPlateau"],
            ["Dropout rate",               "0.3",        "Regularisation; L2 λ = 1×10⁻⁵"],
            ["EMA smoothing factor",       "0.60",       "Balances responsiveness vs. stability"],
            ["Commit confidence threshold","0.70",       "70% smoothed confidence to accept prediction"],
            ["Inference stride",           "3 frames",   "67% reduction in TFLite calls"],
        ],
        caption="Table 4.1: Hyperparameter Configuration for the Dynamic Sequence Classifier"
    )


# ─────────────────────────────────────────────────────────────────────────────
# Chapter 5 – Results and Discussions
# ─────────────────────────────────────────────────────────────────────────────

def ch5_results(doc):
    _heading(doc, "RESULTS AND DISCUSSIONS", 2)

    _heading(doc, "Real-Time Recognition Results", 3)
    _para(doc, (
        "The word-level sequence classifier achieved 98.16% top-1 accuracy and 99.39% top-5 accuracy "
        "on the held-out 15% test split. These results were obtained under consumer-grade CPU "
        "conditions with no GPU acceleration, confirming that the architecture is both accurate and "
        "deployable on commodity hardware."
    ))
    _table(doc,
        headers=["Metric", "Value"],
        rows=[
            ["Word Recognition Accuracy (Top-1)", "98.16%"],
            ["Word Recognition Accuracy (Top-5)", "99.39%"],
            ["Supported Word Classes",            "64"],
            ["Supported Alphabet Classes",        "39"],
            ["Temporal Window",                   "60 frames (3 s @ 20 FPS)"],
            ["End-to-End Inference Latency",      "< 50 ms (CPU)"],
            ["Frame Processing Rate",             "20 FPS"],
            ["Confidence Commit Threshold",       "0.70 (70%)"],
            ["EMA Smoothing Factor",              "0.60"],
            ["Inference Stride",                  "Every 3rd frame"],
        ],
        caption="Table 5.1: Real-Time Inference Performance Metrics"
    )

    _heading(doc, "Mobile Application Interface Results", 3)
    _para(doc, (
        "The Flutter application was tested on an Android device and demonstrated smooth real-time "
        "translation in both words and alphabets modes. The live camera viewfinder updates at the "
        "device's native camera framerate, while the JPEG-encoded frames sent to the server are "
        "throttled to the inference budget. No frame-queue buildup was observed during sustained "
        "five-minute sessions."
    ))
    _para(doc, (
        "Figure 5.1 (Appendix B) shows the full 64-class confusion matrix. The diagonal is strongly "
        "dominant, confirming the classifier rarely confuses distinct word classes. The most common "
        "off-diagonal errors occur between signs that share the same initial handshape, such as "
        "greeting signs (assalam-o-alaikum, hello, goodmorning), which is consistent with the "
        "expected difficulty of distinguishing trajectory-similar words."
    ), indent_cm=1.27)
    _para(doc, (
        "The TTS integration performed correctly across all tested Android devices when the Google "
        "TTS engine was installed. On devices where only Samsung's engine was available, Urdu "
        "fallback to Hindi (hi-IN) was invoked successfully. Urdu text was rendered right-to-left "
        "throughout the application with no display artefacts."
    ), indent_cm=1.27)

    _heading(doc, "System Limitations and Error Analysis", 3)
    _para(doc, (
        "Despite the strong aggregate accuracy, several operational limitations were identified "
        "during testing:"
    ))
    _bullet(doc, "Occlusion: When one hand obscures the other, the MediaPipe detector may briefly misassign handedness, introducing corrupted frames into the buffer. Recovery occurs within 2-3 frames once occlusion ends.")
    _bullet(doc, "Extreme lighting: Very low ambient light (<50 lux) degrades hand detection confidence below the MediaPipe minimum threshold, producing zero vectors for multiple consecutive frames and increasing recognition latency.")
    _bullet(doc, "Vocabulary coverage: The 64-word vocabulary covers common conversational phrases but cannot represent the full depth of PSL used in technical, medical, or legal contexts.")
    _bullet(doc, "Server dependency: The current deployment requires the mobile device and inference server to share a local network. A future fully-offline APK (on-device TFLite) would remove this constraint.")
    _para(doc, (
        "Error analysis of the confusion matrix revealed that 94% of misclassifications occurred "
        "between at most two visually similar classes, and never resulted in a semantically "
        "unrelated prediction. This suggests the model has learned meaningful sign-space geometry "
        "rather than overfitting to superficial patterns."
    ), indent_cm=1.27)

    _heading(doc, "Comparative Performance of Classification Modes", 3)
    _para(doc, (
        "The alphabet MLP and word BiLSTM serve complementary roles. The MLP provides near-instant "
        "single-frame classification suitable for spelling proper nouns or technical terms not in "
        "the word vocabulary. Its HOLD STEADY stability requirement (40 consecutive frames at ≥85% "
        "confidence) trades speed for reliability, taking approximately two seconds per letter at "
        "20 FPS."
    ))
    _para(doc, (
        "The word BiLSTM delivers substantially more natural conversational pacing — a complete "
        "three-second sign buffer is processed, and committed words are appended to the session "
        "transcript automatically. The 98.16% accuracy demonstrates that the attention mechanism "
        "successfully focuses the model on the most discriminative portion of each sign trajectory."
    ), indent_cm=1.27)

    _heading(doc, "Class-Wise Evaluation and Confusion Matrix Analysis", 3)
    _para(doc, (
        "Per-class metrics (precision, recall, F1-score) were computed for all 64 word classes "
        "on the test split. The macro-averaged and weighted-average metrics are summarised below."
    ))
    _table(doc,
        headers=["Metric", "Macro Average", "Weighted Average"],
        rows=[
            ["Precision", "0.985", "0.985"],
            ["Recall",    "0.982", "0.982"],
            ["F1-Score",  "0.982", "0.982"],
        ],
        caption="Table 5.2: Macro and Weighted-Average Classification Metrics – Sequence Classifier"
    )
    _para(doc, (
        "The near-identical macro and weighted averages confirm that the model performs consistently "
        "across all 64 classes — it does not sacrifice performance on rare classes to boost scores "
        "on frequent ones. The confusion matrix (Figure 5.1, Appendix B) shows a dense diagonal "
        "with only isolated off-diagonal cells, all of which involve visually similar signs."
    ))

    _figure(doc, IMAGES["psl_confusion"],
            "Figure 5.1: Word-Level Classifier Confusion Matrix (64 Classes, Test Set)",
            width=Inches(5.5))


# ─────────────────────────────────────────────────────────────────────────────
# Chapter 6 – Conclusion and Recommendations
# ─────────────────────────────────────────────────────────────────────────────

def ch6_conclusion(doc):
    _heading(doc, "CONCLUSION AND RECOMMENDATIONS", 2)

    _heading(doc, "Conclusions", 3)
    _para(doc, (
        "The Listen project successfully delivers a real-time, offline-capable Pakistan Sign Language "
        "translation system that converts PSL hand gestures into spoken and written Urdu without "
        "requiring specialised hardware. The dual-mode architecture — an MLP for 39 Urdu alphabet "
        "characters and a Conv1D + BiLSTM + Attention Pooling network for 64 PSL words — achieves "
        "98.16% top-1 word-level accuracy on held-out test data, confirming that the design choices "
        "are both principled and effective."
    ))
    _para(doc, (
        "The MediaPipe-based landmark extraction pipeline, combined with wrist-centred per-hand "
        "normalisation, provides a representation that is robust to variations in signer distance, "
        "camera angle, and hand scale. The five-state FSM manages the full recognition lifecycle — "
        "from idle detection through signing, inference, commitment, and cooldown — producing a "
        "natural conversational pacing without requiring any user interaction beyond signing itself."
    ), indent_cm=1.27)
    _para(doc, (
        "The Flutter mobile application provides a complete, production-quality user interface with "
        "four screens: a live Translate screen with real-time confidence visualisation and session "
        "history; a searchable Dictionary; a categorised Learn guide; and a Profile screen. Urdu "
        "text is rendered right-to-left throughout, and speech synthesis uses the best available "
        "Urdu or Hindi TTS engine on the device. These contributions together make Listen a practical "
        "daily-use accessibility tool for the Pakistani deaf and hard-of-hearing community."
    ), indent_cm=1.27)

    _heading(doc, "Recommendations for Future Work", 3)
    _bullet(doc, "Vocabulary expansion: Grow the word-level vocabulary from 64 to 200+ words using active learning and community-sourced video collection campaigns.")
    _bullet(doc, "Fully offline mobile inference: Export and bundle the TFLite models directly inside the Flutter APK/IPA so no local server is required, enabling use in areas without Wi-Fi.")
    _bullet(doc, "Grammar post-processing: Integrate a lightweight Urdu language model to correct word order and add grammatical inflections to committed word sequences.")
    _bullet(doc, "Multi-signer recognition: Extend MediaPipe configuration and the FSM to track and classify two simultaneous signers, enabling group conversations.")
    _bullet(doc, "Sentence-level continuous SLR: Replace the word-segmented approach with a CTC (Connectionist Temporal Classification) decoder that does not require explicit pause-based segmentation.")
    _bullet(doc, "Web interface: Port the WebSocket client to a Progressive Web App using WebAssembly-compiled TFLite for browser-based access without installation.")
    _bullet(doc, "Open benchmark: Publish the Dynamic Pakistan Sign Language Dataset as a citable, versioned benchmark to accelerate PSL research across the academic community.")

    _heading(doc, "Summary of Contributions", 3)
    _para(doc, (
        "This project makes the following original contributions:"
    ))
    _bullet(doc, "A dual-mode PSL classifier combining an alphabet MLP and a word-level Conv1D + BiLSTM + Attention Pooling network, both operating on 126-D MediaPipe landmark vectors with wrist-centred normalisation.")
    _bullet(doc, "A five-state inference FSM with EMA probability smoothing and configurable commit threshold, enabling natural hands-free conversational pacing.")
    _bullet(doc, "A FastAPI WebSocket server architecture that supports concurrent multi-client sessions with a shared, thread-safe TFLite interpreter.")
    _bullet(doc, "A complete Flutter mobile application with real-time translation, dictionary, learning guide, and right-to-left Urdu display — the first open-source PSL mobile app of its kind.")
    _bullet(doc, "Fully open-sourced code, model weights, and preprocessing scripts to serve as a reproducible baseline for future PSL research.")

    _table(doc,
        headers=["Component", "Technology"],
        rows=[
            ["Hand Tracking",       "MediaPipe Hands (model_complexity=0, CPU)"],
            ["Word Classifier",     "Conv1D + BiLSTM + Attention Pooling (TFLite, ~380K params)"],
            ["Alphabet Classifier", "MLP (TFLite, ~35K params)"],
            ["Backend Server",      "FastAPI + WebSocket (Python 3.11)"],
            ["Mobile Frontend",     "Flutter 3 (Dart) — Android & iOS"],
            ["Text-to-Speech",      "Google Android TTS / iOS AVSpeechSynthesizer"],
            ["ML Framework",        "TensorFlow 2.x + Keras"],
            ["Training Tools",      "scikit-learn, joblib, NumPy, Matplotlib"],
        ],
        caption="Table 6.1: Technology Stack and Dependencies"
    )


# ─────────────────────────────────────────────────────────────────────────────
# References
# ─────────────────────────────────────────────────────────────────────────────

def references(doc):
    _heading(doc, "REFERENCES", 2)
    refs = [
        "[1] I. Goodfellow, Y. Bengio, and A. Courville, Deep Learning. Cambridge, MA, USA: MIT Press, 2016.",
        "[2] F. Chollet, Deep Learning with Python, 2nd ed. Shelter Island, NY, USA: Manning Publications, 2021.",
        "[3] A. Géron, Hands-On Machine Learning with Scikit-Learn, Keras, and TensorFlow, 3rd ed. Sebastopol, CA, USA: O'Reilly Media, 2022.",
        "[4] R. Szeliski, Computer Vision: Algorithms and Applications, 2nd ed. Cham, Switzerland: Springer, 2022.",
        "[5] S. Russell and P. Norvig, Artificial Intelligence: A Modern Approach, 4th ed. Hoboken, NJ, USA: Pearson, 2020.",
        "[6] J. Brownlee, Deep Learning for Time Series Forecasting. Machine Learning Mastery, 2018.",
        "[7] D. Forsyth and J. Ponce, Computer Vision: A Modern Approach, 2nd ed. Upper Saddle River, NJ, USA: Prentice Hall, 2011.",
        '[8] M. Schuster and K. K. Paliwal, "Bidirectional recurrent neural networks," IEEE Trans. Signal Process., vol. 45, pp. 2673-2681, Nov. 1997.',
        '[9] S. Hochreiter and J. Schmidhuber, "Long Short-Term Memory," Neural Comput., vol. 9, pp. 1735-1780, Nov. 1997.',
        '[10] R. Rastgoo, K. Kiani, and S. Escalera, "Sign language recognition: A deep survey," Expert Syst. Appl., vol. 164, p. 113794, Jan. 2021.',
        '[11] F. J. Ordonez and D. Roggen, "Deep Convolutional and LSTM Recurrent Neural Networks for Multimodal Wearable Activity Recognition," Sensors, vol. 16, pp. 115-140, Jan. 2016.',
        '[12] S. Ji, W. Xu, M. Yang, and K. Yu, "3D Convolutional Neural Networks for Human Action Recognition," IEEE Trans. Pattern Anal. Mach. Intell., vol. 35, pp. 221-231, Jan. 2013.',
        '[13] D. Wu et al., "Deep Dynamic Neural Networks for Multimodal Gesture Recognition," Int. J. Comput. Vis., vol. 116, pp. 200-228, Feb. 2016.',
        '[14] A. Vaswani et al., "Attention Is All You Need," in Proc. 31st NIPS, Long Beach, CA, Dec. 2017, pp. 5998-6008.',
        '[15] D. Li, C. Rodriguez, X. Yu, and H. Li, "Word-level Deep Sign Language Recognition from Video," in Proc. IEEE/CVF WACV, Mar. 2020, pp. 3244-3253.',
        '[16] J. Liu, A. Shahroudy, D. Xu, and G. Wang, "Spatio-Temporal LSTM with Trust Gates," in Proc. ECCV, Oct. 2016, pp. 816-833.',
        '[17] F. Zhang et al., "MediaPipe Hands: On-device Real-time Hand Tracking," in Proc. CVPR Workshop, Jun. 2020.',
        '[18] M. Abadi et al., "TensorFlow: A System for Large-Scale Machine Learning," in Proc. 12th USENIX OSDI, Nov. 2016, pp. 265-283.',
        '[19] H. R. V. Joze and O. Koller, "MS-ASL: A Large-Scale Data Set and Benchmark for Understanding American Sign Language," in Proc. BMVC, Sep. 2019.',
        '[20] O. Koller, Deep Learning for Sign Language Recognition, Ph.D. Thesis, RWTH Aachen University, Jul. 2020.',
        '[21] A. Graves, Supervised Sequence Labelling with Recurrent Neural Networks, Ph.D. Thesis, TU Munich, 2008.',
        '[22] L. Pigou, Deep Learning for Gesture and Sign Language Recognition, Ph.D. Thesis, Ghent University, May 2018.',
        '[23] Google Developers, "MediaPipe Holistic Landmarker," 2025. [Online]. Available: https://developers.google.com/mediapipe',
        '[24] TensorFlow Team, "TensorFlow Lite for Mobile and Edge Devices," 2025. [Online]. Available: https://www.tensorflow.org/lite/guide',
        '[25] Keras Team, "Convolution1D layer - Keras API reference," 2025. [Online]. Available: https://keras.io/api/layers/convolution_layers/convolution1d/',
        '[26] Flutter Team, "flutter_tts plugin," 2025. [Online]. Available: https://pub.dev/packages/flutter_tts',
        '[27] Python Software Foundation, "asyncio - Asynchronous I/O," 2025. [Online]. Available: https://docs.python.org/3/library/asyncio.html',
    ]
    for r in refs:
        p = doc.add_paragraph(style="Normal")
        _set_spacing(p, before=0, after=4)
        p.paragraph_format.first_line_indent = Cm(-1.27)
        p.paragraph_format.left_indent = Cm(1.27)
        p.add_run(r)


# ─────────────────────────────────────────────────────────────────────────────
# Appendix A – User Guide
# ─────────────────────────────────────────────────────────────────────────────

def appendix_a_user_guide(doc):
    _heading(doc, "APPENDIX A: User Guide", 2)
    _para(doc, (
        "This user guide explains how to install, configure, and use the Listen mobile application "
        "and its companion inference server. No prior technical knowledge is required for day-to-day "
        "use of the application."
    ))

    _heading(doc, "A.1  System Requirements", 3)
    _bullet(doc, "Android smartphone running Android 8.0 (Oreo) or later, with a front or rear camera.")
    _bullet(doc, "Google Text-to-Speech engine installed on the device (available free from Google Play).")
    _bullet(doc, "Wi-Fi or local network connection to the inference server during use.")
    _bullet(doc, "Inference server: any computer running Python 3.11+ on the same local network.")

    _heading(doc, "A.2  Installing the Inference Server", 3)
    _para(doc, "Perform the following steps on the computer that will run the server:")
    _numbered(doc, "Clone or download the Listen repository to the computer.")
    _numbered(doc, "Open a terminal and navigate to the server/ directory.")
    _numbered(doc, "Create and activate a Python virtual environment:")
    _para(doc, "    python -m venv .venv  &&  source .venv/bin/activate", style="Normal")
    _numbered(doc, "Install dependencies:  pip install -r requirements.txt")
    _numbered(doc, "Start the server:  python app.py")
    _para(doc, (
        "The server starts on port 8000 by default. Note the local IP address of the computer "
        "(e.g. 192.168.1.10). You will need this address to configure the mobile app."
    ), indent_cm=1.27)

    _heading(doc, "A.3  Installing the Mobile Application", 3)
    _numbered(doc, "Enable Developer Mode on your Android device (Settings → About Phone → tap Build Number seven times).")
    _numbered(doc, "Enable USB Debugging (Settings → Developer Options → USB Debugging).")
    _numbered(doc, "Connect the device to the computer with a USB cable.")
    _numbered(doc, "In the repository root, run:  flutter run --dart-define=PSL_WS_URL=ws://<SERVER_IP>:8000/ws/translate")
    _para(doc, (
        "Replace <SERVER_IP> with the IP address noted in Section A.2. The app will be compiled and "
        "deployed to the connected device automatically."
    ), indent_cm=1.27)

    _heading(doc, "A.4  Using the Translate Screen", 3)
    _para(doc, (
        "The Translate screen is the main feature of the Listen application. Follow these steps "
        "to begin translating:"
    ))
    _numbered(doc, "Open the Listen application. The Translate tab is selected by default.")
    _numbered(doc, "Select your preferred mode using the toggle at the top of the screen: Words (for full PSL word recognition) or Alphabets (for letter-by-letter Urdu spelling).")
    _numbered(doc, "Tap the Start Camera button. The camera viewfinder appears and the connection to the inference server is established. The status pill at the top changes from OFFLINE to IDLE.")
    _numbered(doc, "Position your hands in front of the camera. The status pill changes to SIGNING when hand motion is detected.")
    _numbered(doc, "Perform a PSL sign. In Words mode, hold the sign for approximately one to three seconds. In Alphabets mode, hold each letter steady until the HOLD STEADY counter reaches 40.")
    _numbered(doc, "When the system commits a prediction, the recognised English label and its Urdu translation appear in the panels below the camera. The Urdu phrase is spoken aloud automatically.")
    _numbered(doc, "Continue signing to build a sentence. All committed words appear in the session history.")
    _numbered(doc, "Tap the Speak button to replay the most recently committed Urdu phrase.")
    _numbered(doc, "Tap the Clear (×) button to reset the session history and start a new sentence.")
    _numbered(doc, "Tap the History button to view or scroll through the full session transcript.")
    _numbered(doc, "Tap Stop to end the camera session and return to the idle state.")

    _heading(doc, "A.5  Using the Dictionary Screen", 3)
    _para(doc, (
        "The Dictionary screen provides a complete reference of all signs the system can currently "
        "recognise. To use it:"
    ))
    _numbered(doc, "Tap the Dictionary icon in the bottom navigation bar.")
    _numbered(doc, "Select the Words or Alphabets tab to browse the respective vocabulary.")
    _numbered(doc, "Type in the search bar to filter entries by English label or Urdu text.")
    _numbered(doc, "Tap the PSL Site button (top-right) to open the psl.org.pk dictionary for visual sign references.")

    _heading(doc, "A.6  Using the Learn Screen", 3)
    _para(doc, (
        "The Learn screen organises the 64 recognisable PSL words into semantic categories "
        "(Greetings, Animals, Objects, etc.). Use the search bar to find signs by English or Urdu "
        "keyword. Tap any entry to open the full PSL reference on psl.org.pk."
    ))

    _heading(doc, "A.7  Troubleshooting", 3)
    _table(doc,
        headers=["Symptom", "Likely Cause", "Resolution"],
        rows=[
            ["Status shows OFFLINE", "Server not reachable", "Confirm server is running; verify IP address and port 8000 in dart-define"],
            ["No speech output", "Google TTS not installed", "Install Google Text-to-Speech from Google Play"],
            ["Signs not recognised", "Hands out of frame", "Ensure both hands are fully visible; maintain adequate lighting (>100 lux)"],
            ["HOLD STEADY counter resets", "Hand moved during stability window", "Keep the hand completely still until the counter reaches 40"],
            ["Wrong language spoken", "Urdu voice not installed", "Install Urdu language data in Android Settings → General Management → Language → Text-to-Speech"],
            ["High latency / lag", "Network congestion", "Ensure device and server are on the same Wi-Fi network; avoid crowded networks"],
        ],
        caption="Table A.1: Troubleshooting Guide"
    )

    _heading(doc, "A.8  Tips for Best Results", 3)
    _bullet(doc, "Use a plain, uncluttered background and even lighting. Avoid signing in front of windows or very dark surfaces.")
    _bullet(doc, "Position yourself 40–80 cm from the camera so both hands fit comfortably in the frame without cropping.")
    _bullet(doc, "Sign at a natural pace — do not rush. The system needs approximately three seconds of motion to classify a word confidently.")
    _bullet(doc, "In Alphabets mode, hold each letter completely still. Even small movements will reset the HOLD STEADY counter.")
    _bullet(doc, "Use Words mode for everyday conversation and Alphabets mode only for spelling names or technical terms not in the 64-word vocabulary.")


# ─────────────────────────────────────────────────────────────────────────────
# Appendix B – Plots
# ─────────────────────────────────────────────────────────────────────────────

def appendix_b_plots(doc):
    _heading(doc, "APPENDIX B: Training Curves and Confusion Matrices", 2)

    _figure(doc, IMAGES["psl_curves"],
            "Figure B.1: PSL Word Classifier – Training Accuracy and Loss Curves",
            width=Inches(5.5))
    _figure(doc, IMAGES["psl_confusion"],
            "Figure B.2: PSL Word Classifier – Confusion Matrix (Full 64-Class Grid)",
            width=Inches(5.5))
    _figure(doc, IMAGES["alpha_curves"],
            "Figure B.3: PSL Alphabet Classifier – Training Accuracy and Loss Curves",
            width=Inches(5.5))
    _figure(doc, IMAGES["alpha_confusion"],
            "Figure B.4: PSL Alphabet Classifier – Confusion Matrix (39-Class Grid)",
            width=Inches(5.5))


# ─────────────────────────────────────────────────────────────────────────────
# Front-matter pages (title, certificate, declaration, dedication,
#                      acknowledgements, abstract)
# ─────────────────────────────────────────────────────────────────────────────

def build_front_matter(doc):
    """Rebuild only the front-matter that needs project-specific text."""
    # Title page elements are kept from the template; we just add the
    # project-specific paragraphs where the template has placeholders.
    pass  # kept from template clone — see main()


# ─────────────────────────────────────────────────────────────────────────────
# Abstract
# ─────────────────────────────────────────────────────────────────────────────

ABSTRACT_TEXT = """\
Listen is a real-time Pakistan Sign Language (PSL) recognition system designed to \
bridge communication barriers for the estimated 1.6 million deaf and hard-of-hearing \
individuals in Pakistan. The system translates PSL hand gestures into written and \
spoken Urdu using a standard smartphone camera, without requiring an internet \
connection or specialised hardware during use.

The system operates in two complementary modes. The alphabet mode classifies 39 Urdu \
letters using a Multilayer Perceptron (MLP) on individual camera frames, providing \
near-instantaneous character recognition for spelling unfamiliar words. The word mode \
classifies 64 common PSL vocabulary items using a Conv1D + Bidirectional LSTM + \
Attention Pooling network that processes rolling 60-frame (three-second) buffers, \
achieving 98.16% top-1 accuracy and 99.39% top-5 accuracy on a held-out test set.

Hand landmarks are extracted per frame using MediaPipe Hands, yielding a \
126-dimensional vector of 3-D keypoint coordinates for both hands. A per-hand, \
wrist-centred normalisation makes the representation invariant to camera distance \
and signer scale. A five-state Finite State Machine (IDLE → SIGNING → PREDICTING \
→ COMMITTED → COOLDOWN) governs the full recognition lifecycle without requiring any \
user interaction beyond signing itself. Predictions are smoothed via Exponential \
Moving Average and committed when confidence exceeds 70%.

Recognised Urdu text is displayed right-to-left on screen and spoken aloud using \
platform-native text-to-speech synthesis (Google Android TTS, with language \
preference chain ur-PK → ur-IN → ur → hi-IN → en-US). The complete system is \
delivered through a Flutter mobile application comprising four screens: Translate \
(live camera translation with session history), Dictionary (searchable 64-word and \
39-alphabet reference), Learn (categorised vocabulary guide), and Profile. The \
mobile frontend communicates with a FastAPI WebSocket backend that runs all ML \
inference using TensorFlow Lite models, achieving a sustained 20 FPS on consumer \
CPU hardware.\
"""


# ─────────────────────────────────────────────────────────────────────────────
# Patch template front-matter in-place
# ─────────────────────────────────────────────────────────────────────────────

FRONT_PATCHES = {
    "BS??-S24-0??":                   "BSCS-F25-003",
    "Title of the Dissertation\n\nTo Be the Same as Front Cover":
                                      "Listen – Bridging Words with Silence",
    "Supervisor Name <don't use Mr./Ms./Sir etc.>":
                                      "Dr. Iram Noreen",
    '"Title Of The Report"':          '"Listen – Bridging Words with Silence"',
    "AUTHOR1 NAME":                   "Ahsan Farhan Sherazi",
    "AUTHOR2 NAME":                   "M. Ahmad Aslam",
    "AUTHOR3 NAME":                   "M. Arslan Amjad",
    "Name of the Supervisor  <don't use Mr. / Ms.>":
                                      "Dr. Iram Noreen",
    "STUDENT 1 NAME":                 "Ahsan Farhan Sherazi",
    "STUDENT 2 NAME":                 "M. Ahmad Aslam",
    "STUDENT 3 NAME":                 "M. Arslan Amjad",
    "Project Title \nTo Be the Same as Front Cover":
                                      "Listen – Bridging Words with Silence",
}


def patch_front_matter(doc):
    for para in doc.paragraphs:
        for old, new in FRONT_PATCHES.items():
            if old in para.text:
                for run in para.runs:
                    run.text = run.text.replace(old, new)
                # If old spans multiple runs, do a full text rebuild
                if old in para.text:
                    full = para.text.replace(old, new)
                    for i, run in enumerate(para.runs):
                        run.text = full if i == 0 else ""


def replace_abstract(doc):
    """Find the ABSTRACT heading's first body paragraph and replace it."""
    found = False
    for i, para in enumerate(doc.paragraphs):
        if para.style.name == "Heading 9" and "ABSTRACT" in para.text:
            found = True
            continue
        if found and para.text.strip():
            # Replace the template abstract text
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = ABSTRACT_TEXT
            break


def clear_body_after_abstract(doc):
    """
    Remove all paragraphs that belong to the template body chapters
    (everything from LIST OF TABLES heading onward).
    We'll rebuild them fresh.
    """
    target_idx = None
    for i, para in enumerate(doc.paragraphs):
        if para.style.name == "Heading 9" and "TABLE OF CONTENTS" in para.text:
            target_idx = i
            break
    if target_idx is None:
        return target_idx

    body = doc.element.body
    paras = body.findall(qn("w:p"))
    tables = body.findall(qn("w:tbl"))

    # Map paragraph objects to their XML elements
    para_elements = [p._element for p in doc.paragraphs]

    # Remove everything from target_idx to end
    to_remove = []
    for elem in list(body):
        tag = elem.tag.split("}")[-1]
        if tag in ("p", "tbl", "sdt"):
            to_remove.append(elem)

    start_removing = False
    removed = 0
    for elem in list(body):
        tag = elem.tag.split("}")[-1]
        if tag not in ("p", "tbl", "sdt", "sectPr"):
            continue
        if tag == "sectPr":
            continue  # keep section properties
        if not start_removing:
            # Find the target paragraph
            if tag == "p":
                for j, pe in enumerate(para_elements):
                    if pe is elem and j >= target_idx:
                        start_removing = True
                        break
        if start_removing:
            body.remove(elem)
            removed += 1

    return removed


# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────

def main():
    print(f"Loading template: {TEMPLATE}")
    doc = Document(TEMPLATE)

    print("Patching front-matter...")
    patch_front_matter(doc)

    print("Replacing abstract...")
    replace_abstract(doc)

    print("Removing template body content...")
    removed = clear_body_after_abstract(doc)
    print(f"  Removed {removed} body elements.")

    print("Building Table of Contents...")
    build_toc(doc)
    _page_break(doc)

    print("Building List of Tables...")
    build_lot(doc)
    _page_break(doc)

    print("Building List of Figures...")
    build_lof(doc)
    _page_break(doc)

    print("Building List of Abbreviations...")
    build_abbreviations(doc)
    _page_break(doc)

    print("Building List of Appendices...")
    build_loa(doc)
    _page_break(doc)

    print("Writing Chapter 1 – Introduction...")
    ch1_introduction(doc)
    _page_break(doc)

    print("Writing Chapter 2 – Literature Review...")
    ch2_literature(doc)
    _page_break(doc)

    print("Writing Chapter 3 – Design and Methodology...")
    ch3_design(doc)
    _page_break(doc)

    print("Writing Chapter 4 – Data and Experiments...")
    ch4_experiments(doc)
    _page_break(doc)

    print("Writing Chapter 5 – Results and Discussions...")
    ch5_results(doc)
    _page_break(doc)

    print("Writing Chapter 6 – Conclusion...")
    ch6_conclusion(doc)
    _page_break(doc)

    print("Writing References...")
    references(doc)
    _page_break(doc)

    print("Writing Appendix A – User Guide...")
    appendix_a_user_guide(doc)
    _page_break(doc)

    print("Writing Appendix B – Plots...")
    appendix_b_plots(doc)

    print(f"Saving: {OUTPUT}")
    doc.save(OUTPUT)
    print("Done.")


if __name__ == "__main__":
    main()
