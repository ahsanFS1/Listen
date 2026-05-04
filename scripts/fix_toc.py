"""
Rebuild the TOC, LOT and LOF in fyp document FINAL.docx to reflect
the corrected section numbers from fix_formatting.py.
Run AFTER fix_formatting.py.
"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = DST = "fyp document FINAL.docx"

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _set_line_spacing_15(para):
    pPr = para._element.get_or_add_pPr()
    for sp in pPr.findall(qn("w:spacing")):
        pPr.remove(sp)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:line"), "360")
    sp.set(qn("w:lineRule"), "auto")
    pPr.append(sp)


def _run_font(run, size_pt=12, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")


def _toc_para(doc, num, title, page, style_name, bold=False):
    p = doc.add_paragraph(style=style_name)
    _set_line_spacing_15(p)
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    text = f"{num}\t{title}\t{page}" if num else f"{title}\t{page}"
    r = p.add_run(text)
    _run_font(r, bold=bold)
    return p


def _tof_para(doc, ref, title, page):
    p = doc.add_paragraph(style="table of figures")
    _set_line_spacing_15(p)
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    r = p.add_run(f"{ref}: {title}\t{page}")
    _run_font(r)
    return p


def _heading9(doc, text):
    p = doc.add_paragraph(style="Heading 9")
    _set_line_spacing_15(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    _run_font(r, bold=True)
    return p


def _page_break(doc):
    doc.add_page_break()


# ── Full TOC data ─────────────────────────────────────────────────────────────

TOC_ITEMS = [
    # (num, title, page, style, bold)
    ("",     "DECLARATION",                                      "iii",  "toc 9", True),
    ("",     "ACKNOWLEDGEMENTS",                                 "v",    "toc 9", True),
    ("",     "ABSTRACT",                                         "vii",  "toc 9", True),
    ("",     "TABLE OF CONTENTS",                                "ix",   "toc 9", True),
    ("",     "LIST OF TABLES",                                   "xi",   "toc 9", True),
    ("",     "LIST OF FIGURES",                                  "xiii", "toc 9", True),
    ("",     "LIST OF SYMBOLS / ABBREVIATIONS",                  "xv",   "toc 9", True),
    ("",     "LIST OF APPENDICES",                               "xvi",  "toc 9", True),
    ("",     "CHAPTERS",                                         "",     "toc 9", True),
    ("1",    "INTRODUCTION",                                     "1",    "toc 2", True),
    ("1.1",  "Background",                                       "1",    "toc 3", False),
    ("1.2",  "Problem Statement",                                "2",    "toc 3", False),
    ("1.3",  "Aims and Objectives",                              "2",    "toc 3", False),
    ("1.3.1","Defining the Principles of Real-Time PSL Recognition", "2","toc 4", False),
    ("1.3.2","Implementing the Dual-Mode Classification Architecture","3","toc 4", False),
    ("1.3.3","Integrating the Inference State Machine and Speech Synthesis","3","toc 4",False),
    ("1.4",  "Scope of Project",                                 "4",    "toc 3", False),
    ("2",    "LITERATURE REVIEW",                                "5",    "toc 2", True),
    ("2.1",  "Related Work in Sign Language Recognition",        "5",    "toc 3", False),
    ("2.2",  "Technical Methodologies and Computational Efficiency","6", "toc 3", False),
    ("2.3",  "Dataset Preparation and Preprocessing Pipeline",   "7",    "toc 3", False),
    ("2.4",  "Anatomical Landmark Extraction and Normalisation",  "8",   "toc 3", False),
    ("2.5",  "Temporal Sequence Buffering and Data Stratification","9",  "toc 3", False),
    ("2.6",  "Evaluation Metrics and Model Performance",         "10",   "toc 3", False),
    ("3",    "DESIGN AND METHODOLOGY",                           "12",   "toc 2", True),
    ("3.1",  "System Architecture Overview",                     "12",   "toc 3", False),
    ("3.2",  "Mobile Frontend Architecture",                     "13",   "toc 3", False),
    ("3.2.1","Application Screens",                              "14",   "toc 4", False),
    ("3.2.2","WebSocket Communication Layer",                    "15",   "toc 4", False),
    ("3.2.3","Text-to-Speech Integration",                       "15",   "toc 4", False),
    ("3.3",  "Temporal Sequence Model Design",                   "16",   "toc 3", False),
    ("3.4",  "Real-Time Inference and State Machine Implementation","17", "toc 3", False),
    ("3.4.1","Latency Optimisation",                             "18",   "toc 4", False),
    ("4",    "DATA AND EXPERIMENTS",                             "20",   "toc 2", True),
    ("4.1",  "Data Acquisition and Preprocessing Pipeline",      "20",   "toc 3", False),
    ("4.2",  "Classification Architectures and Implementation",  "21",   "toc 3", False),
    ("4.3",  "Training Methodology and Experimental Setup",      "22",   "toc 3", False),
    ("4.4",  "Hyperparameter Optimisation and Inference Latency","23",   "toc 3", False),
    ("5",    "RESULTS AND DISCUSSIONS",                          "25",   "toc 2", True),
    ("5.1",  "Real-Time Recognition Results",                    "25",   "toc 3", False),
    ("5.2",  "Mobile Application Interface Results",             "26",   "toc 3", False),
    ("5.3",  "System Limitations and Error Analysis",            "27",   "toc 3", False),
    ("5.4",  "Comparative Performance of Classification Modes",  "28",   "toc 3", False),
    ("5.5",  "Class-Wise Evaluation and Confusion Matrix Analysis","29", "toc 3", False),
    ("6",    "CONCLUSION AND RECOMMENDATIONS",                   "31",   "toc 2", True),
    ("6.1",  "Conclusions",                                      "31",   "toc 3", False),
    ("6.2",  "Recommendations for Future Work",                  "32",   "toc 3", False),
    ("6.3",  "Summary of Contributions",                         "33",   "toc 3", False),
    ("",     "REFERENCES",                                       "34",   "toc 2", True),
    ("",     "APPENDICES",                                       "36",   "toc 2", True),
    ("A",    "User Guide",                                       "36",   "toc 3", False),
    ("B",    "Training Curves and Confusion Matrices",           "40",   "toc 3", False),
]

LOT_ITEMS = [
    ("Table 2.1", "Pipeline Processing Times for Different Hardware Configurations", "7"),
    ("Table 2.2", "Evaluation Metrics and Model Complexity for Classification Models", "11"),
    ("Table 3.1", "Inference Latency (ms) Across Primary Pipeline Stages", "19"),
    ("Table 4.1", "Hyperparameter Configuration for the Dynamic Sequence Classifier", "24"),
    ("Table 5.1", "Real-Time Inference Performance Metrics", "25"),
    ("Table 5.2", "Macro and Weighted-Average Classification Metrics - Sequence Classifier", "30"),
    ("Table 6.1", "Technology Stack and Dependencies", "34"),
    ("Table A.1", "Troubleshooting Guide", "39"),
]

LOF_ITEMS = [
    ("Figure 2.1", "Word-Level Classifier Training Curves (Accuracy and Loss)", "11"),
    ("Figure 3.1", "Translate Screen - Idle State", "14"),
    ("Figure 3.2", "PSL Dictionary Screen", "14"),
    ("Figure 4.1", "Alphabet Recognition - SIGNING State (Hold Steady 31/40)", "21"),
    ("Figure 4.2", "Word Recognition - COMMITTED (assalam-o-alaikum, 81.6%)", "21"),
    ("Figure 5.1", "Word-Level Classifier Confusion Matrix (64 Classes, Test Set)", "30"),
    ("Figure B.1", "PSL Word Classifier - Training Accuracy and Loss Curves", "41"),
    ("Figure B.2", "PSL Word Classifier - Confusion Matrix (Full 64-Class Grid)", "42"),
    ("Figure B.3", "PSL Alphabet Classifier - Training Accuracy and Loss Curves", "43"),
    ("Figure B.4", "PSL Alphabet Classifier - Confusion Matrix (39-Class Grid)", "44"),
]


def find_section_range(doc, start_heading, end_heading=None):
    """Return (start_idx, end_idx) of paragraphs between two Heading 9 markers."""
    paras = doc.paragraphs
    start = end = None
    for i, p in enumerate(paras):
        if p.style.name == "Heading 9" and start_heading in p.text:
            start = i
        elif start is not None and end_heading and p.style.name == "Heading 9" and end_heading in p.text:
            end = i
            break
    if start is not None and end is None:
        end = len(paras)
    return start, end


def remove_paragraphs_between(doc, start_idx, end_idx):
    """Remove all paragraphs between start_idx (exclusive) and end_idx (exclusive)."""
    body = doc.element.body
    paras = doc.paragraphs
    to_remove = [paras[i]._element for i in range(start_idx + 1, end_idx)]
    for elem in to_remove:
        if elem.getparent() is not None:
            elem.getparent().remove(elem)


def insert_after(doc, anchor_para, new_paras_fn):
    """Insert paragraphs returned by new_paras_fn() after anchor_para."""
    # We rebuild via a temp doc then graft elements
    tmp = Document()
    new_paras_fn(tmp)
    anchor_elem = anchor_para._element
    parent = anchor_elem.getparent()
    # Insert in reverse so they come out in order
    tmp_paras = list(tmp.element.body)
    # Remove the default empty paragraph that Document() adds
    insert_idx = list(parent).index(anchor_elem) + 1
    for elem in tmp_paras:
        tag = elem.tag.split("}")[-1]
        if tag in ("p", "tbl"):
            import copy
            parent.insert(insert_idx, copy.deepcopy(elem))
            insert_idx += 1


def rebuild_list_section(doc, heading_text, next_heading_text, build_fn):
    """Find the section, remove its content, rebuild it."""
    paras = doc.paragraphs
    start_idx = end_idx = None
    for i, p in enumerate(paras):
        if p.style.name == "Heading 9" and heading_text in p.text:
            start_idx = i
        elif start_idx is not None and p.style.name == "Heading 9" and next_heading_text in p.text:
            end_idx = i
            break
    if start_idx is None:
        print(f"  [WARN] Could not find section: {heading_text}")
        return
    if end_idx is None:
        end_idx = len(paras)

    # Remove everything between start and end (exclusive of both)
    body = doc.element.body
    all_paras = list(body)  # includes sectPr, page breaks etc
    # Identify the XML elements
    heading_elem = paras[start_idx]._element
    end_elem = paras[end_idx]._element if end_idx < len(paras) else None

    # Collect elements to delete: between heading_elem and end_elem
    collecting = False
    to_delete = []
    for elem in list(body):
        if elem is heading_elem:
            collecting = True
            continue
        if end_elem is not None and elem is end_elem:
            break
        if collecting:
            tag = elem.tag.split("}")[-1]
            if tag in ("p", "tbl", "sdt"):
                to_delete.append(elem)

    for e in to_delete:
        body.remove(e)

    # Now insert new content after heading_elem
    import copy
    tmp = Document()
    # Copy styles from doc to tmp (we just need to write paragraphs with plain text)
    build_fn(tmp)

    insert_idx = list(body).index(heading_elem) + 1
    for elem in list(tmp.element.body):
        tag = elem.tag.split("}")[-1]
        if tag in ("p", "tbl"):
            body.insert(insert_idx, copy.deepcopy(elem))
            insert_idx += 1


def build_toc_content(tmp):
    for num, title, page, style_name, bold in TOC_ITEMS:
        try:
            p = tmp.add_paragraph(style=style_name)
        except Exception:
            p = tmp.add_paragraph()
        _set_line_spacing_15(p)
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after  = Pt(0)
        text = f"{num}\t{title}\t{page}" if num else (f"{title}\t{page}" if page else title)
        r = p.add_run(text)
        _run_font(r, bold=bold)


def build_lot_content(tmp):
    for ref, title, page in LOT_ITEMS:
        try:
            p = tmp.add_paragraph(style="table of figures")
        except Exception:
            p = tmp.add_paragraph()
        _set_line_spacing_15(p)
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after  = Pt(0)
        r = p.add_run(f"{ref}: {title}\t{page}")
        _run_font(r)


def build_lof_content(tmp):
    for ref, title, page in LOF_ITEMS:
        try:
            p = tmp.add_paragraph(style="table of figures")
        except Exception:
            p = tmp.add_paragraph()
        _set_line_spacing_15(p)
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after  = Pt(0)
        r = p.add_run(f"{ref}: {title}\t{page}")
        _run_font(r)


def main():
    print(f"Loading: {SRC}")
    doc = Document(SRC)

    print("Rebuilding TABLE OF CONTENTS...")
    rebuild_list_section(doc, "TABLE OF CONTENTS", "LIST OF TABLES", build_toc_content)

    print("Rebuilding LIST OF TABLES...")
    rebuild_list_section(doc, "LIST OF TABLES", "LIST OF FIGURES", build_lot_content)

    print("Rebuilding LIST OF FIGURES...")
    rebuild_list_section(doc, "LIST OF FIGURES", "LIST OF SYMBOLS", build_lof_content)

    print(f"Saving: {DST}")
    doc.save(DST)
    print("Done.")


if __name__ == "__main__":
    main()
