"""
Post-processor for fyp document FINAL.docx

Fixes two things:
  1. Section numbering  – prepends "1.1 ", "2.3 ", "3.2.1 " etc. to every
     Heading 3 / Heading 4 paragraph in the chapter body (front-matter and
     Appendix headings that already carry numbers / letters are left alone).
  2. Formatting        – makes every Normal paragraph match the template:
       - 12 pt Times New Roman
       - 1.5 line spacing (auto, 360 twips)
       - Justified alignment
       - No explicit space_before / space_after (template relies on line spacing)
       - First-line indent 1.27 cm for body paragraphs (after the first para
         in each section), 0 for the first paragraph.
     Heading 2 paragraphs get the correct chapter-title treatment:
       - Bold, centered, 12 pt, 1.5 line spacing, all caps.
     Heading 3 / 4 paragraphs get:
       - Bold, justified, 12 pt, 1.5 line spacing.
     Heading 9 (front-matter titles like ABSTRACT, TABLE OF CONTENTS) stay
     centered and bold – untouched.

Run from repo root:
    python scripts/fix_formatting.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as ET
import copy

SRC = "fyp document FINAL.docx"
DST = "fyp document FINAL.docx"

# ── Heading number map ────────────────────────────────────────────────────────
# Maps (chapter_index, h3_counter, h4_counter) → number prefix string.
# We derive this by scanning the document headings in order.

# Chapter titles (Heading 2) that mark the start of a numbered chapter.
NUMBERED_CHAPTERS = {
    "INTRODUCTION",
    "LITERATURE REVIEW",
    "DESIGN AND METHODOLOGY",
    "DATA AND EXPERIMENTS",
    "RESULTS AND DISCUSSIONS",
    "CONCLUSION AND RECOMMENDATIONS",
}

# Heading 2 titles that are NOT numbered chapters.
UNNUMBERED_H2 = {"REFERENCES", "APPENDIX A: User Guide",
                 "APPENDIX B: Training Curves and Confusion Matrices"}

# Front-matter Heading 9 set (leave alone).
FRONT_MATTER_H9 = {
    "DECLARATION", "ACKNOWLEDGEMENTS", "ABSTRACT",
    "TABLE OF CONTENTS", "LIST OF TABLES", "LIST OF FIGURES",
    "LIST OF SYMBOLS / ABBREVIATIONS", "LIST OF APPENDICES",
}

# Appendix subsections already have "A.x" in their text — skip re-numbering.
def _already_numbered(text):
    t = text.strip()
    # "A.1", "A.2", ..., "A.8", "B" or starts with digit like "1.1"
    if t and (t[0].isdigit() or (len(t) > 1 and t[0].isalpha() and t[1] == '.')):
        return True
    return False


# ── XML helpers ───────────────────────────────────────────────────────────────

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _set_line_spacing_15(para):
    """Set 1.5 line spacing (auto, 360 twips) on a paragraph element."""
    pPr = para._element.get_or_add_pPr()
    # Remove existing spacing element then set fresh
    for sp in pPr.findall(qn("w:spacing")):
        pPr.remove(sp)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:line"), "360")
    sp.set(qn("w:lineRule"), "auto")
    # Remove before/after to match template (template has none)
    pPr.append(sp)


def _set_alignment(para, alignment):
    pf = para.paragraph_format
    pf.alignment = alignment


def _set_first_line_indent(para, cm):
    pf = para.paragraph_format
    if cm == 0:
        pf.first_line_indent = None
        pf.left_indent = None
    else:
        pf.first_line_indent = Cm(cm)


def _clear_space_before_after(para):
    pPr = para._element.get_or_add_pPr()
    for sp in pPr.findall(qn("w:spacing")):
        sp.attrib.pop(qn("w:before"), None)
        sp.attrib.pop(qn("w:after"), None)
        sp.attrib.pop(qn("w:beforeLines"), None)
        sp.attrib.pop(qn("w:afterLines"), None)


def _set_run_font(run, size_pt=12, bold=None, name="Times New Roman"):
    run.font.name = name
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    # East-Asia font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)


def _format_normal(para, first_in_section=False):
    """Apply template body-text formatting to a Normal paragraph."""
    _set_alignment(para, WD_ALIGN_PARAGRAPH.JUSTIFY)
    _set_line_spacing_15(para)
    _clear_space_before_after(para)
    # First paragraph in a section has no indent; subsequent ones do.
    if first_in_section:
        _set_first_line_indent(para, 0)
    else:
        _set_first_line_indent(para, 1.27)
    for run in para.runs:
        _set_run_font(run, size_pt=12, bold=None)


def _format_h2(para):
    """Chapter title: bold, centered, 12pt, 1.5 line spacing."""
    _set_alignment(para, WD_ALIGN_PARAGRAPH.CENTER)
    _set_line_spacing_15(para)
    _clear_space_before_after(para)
    _set_first_line_indent(para, 0)
    for run in para.runs:
        _set_run_font(run, size_pt=12, bold=True)


def _format_h3(para):
    """Section heading: bold, justified, 12pt, 1.5 line spacing."""
    _set_alignment(para, WD_ALIGN_PARAGRAPH.JUSTIFY)
    _set_line_spacing_15(para)
    _clear_space_before_after(para)
    _set_first_line_indent(para, 0)
    for run in para.runs:
        _set_run_font(run, size_pt=12, bold=True)


def _format_h4(para):
    """Sub-section heading: bold, justified, 12pt, 1.5 line spacing."""
    _set_alignment(para, WD_ALIGN_PARAGRAPH.JUSTIFY)
    _set_line_spacing_15(para)
    _clear_space_before_after(para)
    _set_first_line_indent(para, 0)
    for run in para.runs:
        _set_run_font(run, size_pt=12, bold=True)


def _format_caption(para):
    """Caption: centered, 12pt, 1.5 line spacing."""
    _set_alignment(para, WD_ALIGN_PARAGRAPH.CENTER)
    _set_line_spacing_15(para)
    _clear_space_before_after(para)
    _set_first_line_indent(para, 0)
    for run in para.runs:
        _set_run_font(run, size_pt=12, bold=False)


def _format_bullet(para):
    """Bullet paragraph: justified, 12pt, 1.5 line spacing."""
    _set_alignment(para, WD_ALIGN_PARAGRAPH.JUSTIFY)
    _set_line_spacing_15(para)
    _clear_space_before_after(para)
    for run in para.runs:
        _set_run_font(run, size_pt=12)


def _format_toc(para):
    """TOC entry: keep as-is but ensure 12pt font."""
    _set_line_spacing_15(para)
    _clear_space_before_after(para)
    for run in para.runs:
        _set_run_font(run, size_pt=12)


def _prepend_number(para, prefix):
    """Prepend 'prefix ' to the paragraph's text, preserving run formatting."""
    if not para.runs:
        para.add_run(prefix + " ")
        return
    first_run = para.runs[0]
    first_run.text = prefix + "\t" + first_run.text


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    print(f"Loading: {SRC}")
    doc = Document(SRC)

    # ── Pass 1: Assign section numbers ────────────────────────────────────────
    chapter = 0          # 1-based chapter counter (only for NUMBERED_CHAPTERS)
    h3_cnt  = 0          # section counter within chapter
    h3_last = 0          # last h3 number (for h4 subsection tracking)
    h4_cnt  = 0          # sub-section counter within h3

    in_numbered_chapter = False
    in_appendix = False

    for para in doc.paragraphs:
        style = para.style.name
        text  = para.text.strip()

        if style == "Heading 2":
            if text in NUMBERED_CHAPTERS:
                chapter += 1
                h3_cnt  = 0
                h4_cnt  = 0
                h3_last = 0
                in_numbered_chapter = True
                in_appendix = False
            elif "APPENDIX" in text:
                in_numbered_chapter = False
                in_appendix = True
            else:
                in_numbered_chapter = False
                in_appendix = False

        elif style == "Heading 3":
            if in_numbered_chapter and not _already_numbered(text):
                h3_cnt += 1
                h4_cnt  = 0
                h3_last = h3_cnt
                prefix = f"{chapter}.{h3_cnt}"
                _prepend_number(para, prefix)
            # Appendix H3s already have A.x in text — skip

        elif style == "Heading 4":
            if in_numbered_chapter and not _already_numbered(text):
                h4_cnt += 1
                prefix = f"{chapter}.{h3_last}.{h4_cnt}"
                _prepend_number(para, prefix)

    # ── Pass 2: Fix formatting ────────────────────────────────────────────────
    # Track "first paragraph after heading" for indent control
    prev_was_heading = True   # start True so first body para gets no indent

    for para in doc.paragraphs:
        style = para.style.name

        if style == "Heading 9":
            # Front-matter titles — just ensure 12pt bold centered
            _set_alignment(para, WD_ALIGN_PARAGRAPH.CENTER)
            _set_line_spacing_15(para)
            _clear_space_before_after(para)
            for run in para.runs:
                _set_run_font(run, size_pt=12, bold=True)
            prev_was_heading = True

        elif style == "Heading 2":
            _format_h2(para)
            prev_was_heading = True

        elif style == "Heading 3":
            _format_h3(para)
            prev_was_heading = True

        elif style == "Heading 4":
            _format_h4(para)
            prev_was_heading = True

        elif style == "Caption":
            _format_caption(para)
            prev_was_heading = False

        elif style in ("toc 2", "toc 3", "toc 4", "toc 9", "table of figures"):
            _format_toc(para)
            prev_was_heading = False

        elif style == "Normal":
            if not para.text.strip():
                # Empty spacer paragraph — keep but set spacing
                _set_line_spacing_15(para)
                _clear_space_before_after(para)
                # Don't flip prev_was_heading for blank lines
                continue
            _format_normal(para, first_in_section=prev_was_heading)
            prev_was_heading = False

        elif style == "List Number":
            _format_bullet(para)
            prev_was_heading = False

        else:
            # Any other style — apply baseline 12pt 1.5 spacing
            _set_line_spacing_15(para)
            _clear_space_before_after(para)
            for run in para.runs:
                _set_run_font(run, size_pt=12)
            prev_was_heading = False

    # ── Pass 3: Fix page margins to match template exactly ────────────────────
    for section in doc.sections:
        section.left_margin   = Cm(4.0)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    # ── Pass 4: Fix table cell text formatting ────────────────────────────────
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _set_line_spacing_15(para)
                    _clear_space_before_after(para)
                    _set_first_line_indent(para, 0)
                    for run in para.runs:
                        _set_run_font(run, size_pt=11)

    print(f"Saving: {DST}")
    doc.save(DST)

    # ── Verify ────────────────────────────────────────────────────────────────
    doc2 = Document(DST)
    headings = [(p.style.name, p.text[:70]) for p in doc2.paragraphs
                if p.style.name in ("Heading 2", "Heading 3", "Heading 4")
                and p.text.strip()]
    print("\n=== Heading structure after fix ===")
    for style, text in headings:
        lvl = {"Heading 2": "CH", "Heading 3": " S", "Heading 4": "  ss"}[style]
        print(f"  {lvl}  {text}")
    print("\nDone.")


if __name__ == "__main__":
    main()
